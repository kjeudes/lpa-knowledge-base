# -*- coding: utf-8 -*-
# ============================================================================
# LPA — MODÈLE AUTONOME v3.1
# Leçon 12 — ANGLAIS LV1 · 3ème — Fashion show (Défilé de mode)
# Professeur : Evelyne ASSI
# DPFC/MENET-FP CI · 2025-2026
# ============================================================================

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch
import numpy as np
import os

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ============================================================
# PALETTE
# ============================================================
BLEU        = "184082"
BLEU_BG     = "EAF2F8"
BLEU_MED    = "D9E2F3"
BLEU_FONCE  = "183E82"
OR          = "C99225"
VERT        = "2A7E5B"
VERT_BG     = "E2F0D9"
JAUNE_BG    = "FFF2CC"
VIOLET      = "6B3FA0"
VIOLET_BG   = "F0EBF8"
ORANGE      = "C0541A"
ORANGE_BG   = "FEF0E7"
GRIS        = "555555"
TEXTE       = "232323"

# ============================================================
# SECTION 1 — HELPERS (NE PAS MODIFIER)
# ============================================================

def _set_cell_bg(cell, hexcolor):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hexcolor)
    cell._tc.get_or_add_tcPr().append(shd)

def _set_cell_borders(cell, color="CCCCCC", sz=4,
                       sides=("top","left","bottom","right"), style="single"):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for side in ("top","left","bottom","right"):
        el = OxmlElement(f"w:{side}")
        if side in sides:
            el.set(qn("w:val"), style)
            el.set(qn("w:sz"), str(sz))
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), color)
        else:
            el.set(qn("w:val"), "nil")
        borders.append(el)
    tcPr.append(borders)

def _set_cell_margins(cell, top=80, bottom=80, left=140, right=140):
    tcPr = cell._tc.get_or_add_tcPr()
    m = OxmlElement("w:tcMar")
    for side, val in (("top",top),("bottom",bottom),("left",left),("right",right)):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        m.append(el)
    tcPr.append(m)

def _full_width_1cell(doc):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    t.columns[0].width = Cm(16.5)
    cell = t.cell(0,0); cell.width = Cm(16.5)
    _set_cell_margins(cell)
    tblPr = t._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for s in ("top","left","bottom","right","insideH","insideV"):
        e = OxmlElement(f"w:{s}"); e.set(qn("w:val"), "nil"); borders.append(e)
    tblPr.append(borders)
    return t, cell

def _para_border(p, color, side="bottom", sz=8):
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    el = OxmlElement(f"w:{side}")
    el.set(qn("w:val"), "single"); el.set(qn("w:sz"), str(sz))
    el.set(qn("w:space"), "1");   el.set(qn("w:color"), color)
    pbdr.append(el); pPr.append(pbdr)

# -------------------------------------------------------
# Runs (texte stylé)
# -------------------------------------------------------
def r(p, txt, bold=False, italic=False, underline=False,
      color=None, size=None, font="Arial"):
    run = p.add_run(txt)
    run.bold = bold; run.italic = italic; run.underline = underline
    run.font.name = font
    if size:  run.font.size = Pt(size)
    if color: run.font.color.rgb = RGBColor.from_string(color)
    return run

def rb(p, txt, color=TEXTE, size=11):
    return r(p, txt, bold=True, color=color, size=size)

def ri(p, txt, color=GRIS, size=11):
    return r(p, txt, italic=True, color=color, size=size)

def ru(p, txt, color=TEXTE, size=11):
    return r(p, txt, underline=True, color=color, size=size)

def rbi(p, txt, color=TEXTE, size=11):
    return r(p, txt, bold=True, italic=True, color=color, size=size)

# -------------------------------------------------------
# Éléments structurels
# -------------------------------------------------------
def add_part_banner(doc, text):
    """Bandeau de section (bleu foncé, texte blanc)."""
    t, cell = _full_width_1cell(doc)
    _set_cell_bg(cell, BLEU_FONCE)
    _set_cell_margins(cell, top=120, bottom=120, left=180)
    p = cell.paragraphs[0]; p.text = ""
    r(p, text, bold=True, color="FFFFFF", size=13)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def add_phase_title(doc, text):
    """Titre de phase (bleu, gras, taille 13)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(4)
    r(p, text, bold=True, color=BLEU, size=13)

def add_subsection(doc, text):
    """Sous-titre (or/amber, gras, taille 11.5)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    r(p, text, bold=True, color=OR, size=11.5)

def add_body(doc, *runs):
    """Paragraphe de corps avec runs mixtes."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(4)
    for item in runs:
        if isinstance(item, str):
            r(p, item, color=TEXTE, size=11)
        else:
            txt, style = item
            bold      = "b" in style
            italic    = "i" in style
            underline = "u" in style
            r(p, txt, bold=bold, italic=italic, underline=underline,
              color=TEXTE, size=11)
    return p

def add_bullet(doc, *runs):
    """Puce. Mêmes arguments que add_body."""
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    for item in runs:
        if isinstance(item, str):
            r(p, item, color=TEXTE, size=11)
        else:
            txt, style = item
            bold      = "b" in style
            italic    = "i" in style
            underline = "u" in style
            r(p, txt, bold=bold, italic=italic, underline=underline,
              color=TEXTE, size=11)

def add_numbered(doc, *runs):
    """Liste numérotée. Mêmes arguments que add_body."""
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(2)
    for item in runs:
        if isinstance(item, str):
            r(p, item, color=TEXTE, size=11)
        else:
            txt, style = item
            bold      = "b" in style
            italic    = "i" in style
            underline = "u" in style
            r(p, txt, bold=bold, italic=italic, underline=underline,
              color=TEXTE, size=11)

# -------------------------------------------------------
# Encadrés existants
# -------------------------------------------------------
def add_harkness(doc, question, guide, corrige_txt):
    """Bloc Harkness — encadré bleu clair."""
    t, cell = _full_width_1cell(doc)
    _set_cell_bg(cell, BLEU_BG)
    _set_cell_borders(cell, BLEU, sz=6, sides=("top","left","bottom","right"))
    _set_cell_margins(cell, top=100, bottom=100, left=160, right=160)
    p = cell.paragraphs[0]; p.text = ""
    r(p, "▶ HARKNESS", bold=True, color=OR, size=9.5)
    for label, txt in [("Question : ", question),
                        ("Guide : ",    guide),
                        ("Corrigé : ",  corrige_txt)]:
        pp = cell.add_paragraph()
        pp.paragraph_format.space_after = Pt(2)
        r(pp, label, bold=True, color=BLEU, size=10.5)
        r(pp, txt,   color=TEXTE, size=10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def add_retenir(doc, title, lines):
    """Encadré 'À retenir' — fond vert clair."""
    t, cell = _full_width_1cell(doc)
    _set_cell_bg(cell, VERT_BG)
    _set_cell_borders(cell, VERT, sz=6, sides=("top","left","bottom","right"))
    _set_cell_margins(cell, top=100, bottom=100, left=160, right=160)
    p = cell.paragraphs[0]; p.text = ""
    r(p, "┌─ " + title + " ─", bold=True, color=VERT, size=9.5)
    for ln in lines:
        pp = cell.add_paragraph(); pp.paragraph_format.space_after = Pt(2)
        if isinstance(ln, str):
            r(pp, ln, color=TEXTE, size=10.5)
        else:
            for item in ln:
                if isinstance(item, str):
                    r(pp, item, color=TEXTE, size=10.5)
                else:
                    txt, style = item
                    r(pp, txt, bold=("b" in style), italic=("i" in style),
                      underline=("u" in style), color=TEXTE, size=10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def add_attention(doc, title, lines):
    """Encadré 'Attention' — fond jaune."""
    t, cell = _full_width_1cell(doc)
    _set_cell_bg(cell, JAUNE_BG)
    _set_cell_borders(cell, OR, sz=6, sides=("top","left","bottom","right"))
    _set_cell_margins(cell, top=100, bottom=100, left=160, right=160)
    p = cell.paragraphs[0]; p.text = ""
    r(p, "⚠  " + title, bold=True, color=OR, size=9.5)
    for ln in lines:
        pp = cell.add_paragraph(); pp.paragraph_format.space_after = Pt(2)
        r(pp, ln, color=TEXTE, size=10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def add_encadre_bleu(doc, title, lines):
    """Encadré bleu clair générique (définitions DPFC, prononciation…)."""
    t, cell = _full_width_1cell(doc)
    _set_cell_bg(cell, BLEU_BG)
    _set_cell_borders(cell, BLEU, sz=4, sides=("top","left","bottom","right"))
    _set_cell_margins(cell, top=100, bottom=100, left=160, right=160)
    p = cell.paragraphs[0]; p.text = ""
    r(p, "◆  " + title, bold=True, color=BLEU, size=9.5)
    for ln in lines:
        pp = cell.add_paragraph(); pp.paragraph_format.space_after = Pt(2)
        if isinstance(ln, str):
            r(pp, ln, color=TEXTE, size=10.5)
        else:
            for item in ln:
                if isinstance(item, str):
                    r(pp, item, color=TEXTE, size=10.5)
                else:
                    txt, style = item
                    r(pp, txt, bold=("b" in style), italic=("i" in style),
                      underline=("u" in style), color=TEXTE, size=10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

# -------------------------------------------------------
# NOUVELLES FONCTIONS — Trace écrite, Analogie CI, Synthèse, Capsule
# -------------------------------------------------------

def add_trace_ecrite(doc, lines):
    """
    Encadré 'Trace écrite' — fond bleu très clair, bordure gauche épaisse bleue.
    lines : liste de str ou de listes de tuples (txt, style).
    Première ligne peut être un titre de section si elle commence par '==='.
    """
    t, cell = _full_width_1cell(doc)
    _set_cell_bg(cell, "EDF4FB")
    _set_cell_borders(cell, BLEU, sz=18, sides=("left",))
    _set_cell_borders(cell, BLEU_MED, sz=4, sides=("top","bottom","right"))
    _set_cell_margins(cell, top=100, bottom=100, left=180, right=160)
    p = cell.paragraphs[0]; p.text = ""
    r(p, "✎  TRACE ÉCRITE  (à recopier dans le cahier)", bold=True, color=BLEU, size=9.5)
    for ln in lines:
        pp = cell.add_paragraph(); pp.paragraph_format.space_after = Pt(2)
        if isinstance(ln, str):
            # Titres internes signalés par "==="
            if ln.startswith("==="):
                r(pp, ln.replace("===","").strip(), bold=True, color=BLEU_FONCE, size=10.5)
            else:
                r(pp, ln, color=TEXTE, size=10.5)
        else:
            for item in ln:
                if isinstance(item, str):
                    r(pp, item, color=TEXTE, size=10.5)
                else:
                    txt, style = item
                    r(pp, txt, bold=("b" in style), italic=("i" in style),
                      underline=("u" in style), color=TEXTE, size=10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_analogie(doc, lines):
    """
    Encadré 'Analogie CI' — fond orange très clair, icône ~, bordure orange.
    lines : liste de str. Chaque str = une ligne d'analogie.
    """
    t, cell = _full_width_1cell(doc)
    _set_cell_bg(cell, ORANGE_BG)
    _set_cell_borders(cell, ORANGE, sz=6, sides=("top","left","bottom","right"))
    _set_cell_margins(cell, top=100, bottom=100, left=160, right=160)
    p = cell.paragraphs[0]; p.text = ""
    r(p, "~  ANALOGIE CI", bold=True, color=ORANGE, size=9.5)
    for ln in lines:
        pp = cell.add_paragraph(); pp.paragraph_format.space_after = Pt(2)
        r(pp, ln, color=TEXTE, size=10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_synthese(doc, points, mots_cles, erreurs):
    """
    Bloc Synthèse — 3 sous-blocs : 5 points essentiels, 5 mots clés, 3 erreurs.
    points     : liste de str (5 éléments)
    mots_cles  : liste de tuples (mot, définition courte) (5 éléments)
    erreurs    : liste de tuples (erreur, correction) (3 éléments)
    """
    t, cell = _full_width_1cell(doc)
    _set_cell_bg(cell, BLEU_MED)
    _set_cell_borders(cell, BLEU, sz=6, sides=("top","left","bottom","right"))
    _set_cell_margins(cell, top=100, bottom=100, left=160, right=160)
    p = cell.paragraphs[0]; p.text = ""
    r(p, "SYNTHÈSE DE LA LEÇON", bold=True, color=BLEU_FONCE, size=11)

    # 5 points essentiels
    pp = cell.add_paragraph(); pp.paragraph_format.space_before = Pt(4)
    r(pp, "5 points essentiels à retenir", bold=True, color=BLEU, size=10.5)
    for i, pt in enumerate(points, 1):
        pp2 = cell.add_paragraph(); pp2.paragraph_format.space_after = Pt(1)
        r(pp2, f"  {i}. ", bold=True, color=BLEU, size=10.5)
        r(pp2, pt, color=TEXTE, size=10.5)

    # 5 mots clés
    pp = cell.add_paragraph(); pp.paragraph_format.space_before = Pt(6)
    r(pp, "5 mots clés", bold=True, color=BLEU, size=10.5)
    for mot, defi in mots_cles:
        pp2 = cell.add_paragraph(); pp2.paragraph_format.space_after = Pt(1)
        r(pp2, f"  · ", bold=True, color=OR, size=10.5)
        r(pp2, mot, bold=True, color=TEXTE, size=10.5)
        r(pp2, f" = {defi}", color=GRIS, size=10.5)

    # 3 erreurs fréquentes
    pp = cell.add_paragraph(); pp.paragraph_format.space_before = Pt(6)
    r(pp, "3 erreurs fréquentes + correction", bold=True, color=BLEU, size=10.5)
    for err, ok in erreurs:
        pp2 = cell.add_paragraph(); pp2.paragraph_format.space_after = Pt(2)
        r(pp2, "  ✗ ", bold=True, color="A62D2D", size=10.5)
        r(pp2, err + "  →  ", italic=True, color="A62D2D", size=10.5)
        r(pp2, "✓ ", bold=True, color=VERT, size=10.5)
        r(pp2, ok, bold=True, color=VERT, size=10.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_capsule_reprise(doc, points_5, questions_3, lien):
    """
    Bloc Capsule de reprise — fond violet très clair.
    points_5   : liste de str (5 points clés)
    questions_3: liste de tuples (question, réponse) (3 éléments)
    lien       : str — lien explicite avec la leçon du jour
    """
    t, cell = _full_width_1cell(doc)
    _set_cell_bg(cell, VIOLET_BG)
    _set_cell_borders(cell, VIOLET, sz=6, sides=("top","left","bottom","right"))
    _set_cell_margins(cell, top=100, bottom=100, left=160, right=160)
    p = cell.paragraphs[0]; p.text = ""
    r(p, "CAPSULE DE REPRISE — Leçon précédente", bold=True, color=VIOLET, size=9.5)

    # 5 points
    pp = cell.add_paragraph(); pp.paragraph_format.space_before = Pt(4)
    r(pp, "5 points clés :", bold=True, color=VIOLET, size=10.5)
    for i, pt in enumerate(points_5, 1):
        pp2 = cell.add_paragraph(); pp2.paragraph_format.space_after = Pt(1)
        r(pp2, f"  Point {i} · ", bold=True, color=VIOLET, size=10.5)
        r(pp2, pt, color=TEXTE, size=10.5)

    # 3 questions flash
    pp = cell.add_paragraph(); pp.paragraph_format.space_before = Pt(6)
    r(pp, "3 questions flash :", bold=True, color=VIOLET, size=10.5)
    for i, (q, rep) in enumerate(questions_3, 1):
        pp2 = cell.add_paragraph(); pp2.paragraph_format.space_after = Pt(1)
        r(pp2, f"  Q{i} · ", bold=True, color=VIOLET, size=10.5)
        r(pp2, q, color=TEXTE, size=10.5)
        r(pp2, "  → ", bold=True, color=OR, size=10.5)
        r(pp2, rep, italic=True, color=TEXTE, size=10.5)

    # Lien avec la leçon du jour
    pp = cell.add_paragraph(); pp.paragraph_format.space_before = Pt(6)
    r(pp, "Lien avec la leçon du jour : ", bold=True, color=VIOLET, size=10.5)
    r(pp, lien, color=TEXTE, size=10.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_ancrage_ivoirien(doc, lines):
    """
    Bloc Ancrage ivoirien — fond vert très clair, bordure gauche verte épaisse.
    lines : liste de str.
    """
    t, cell = _full_width_1cell(doc)
    _set_cell_bg(cell, VERT_BG)
    _set_cell_borders(cell, VERT, sz=18, sides=("left",))
    _set_cell_borders(cell, VERT_BG, sz=4, sides=("top","bottom","right"))
    _set_cell_margins(cell, top=100, bottom=100, left=180, right=160)
    p = cell.paragraphs[0]; p.text = ""
    r(p, "ANCRAGE IVOIRIEN", bold=True, color=VERT, size=9.5)
    for ln in lines:
        pp = cell.add_paragraph(); pp.paragraph_format.space_after = Pt(2)
        r(pp, ln, color=TEXTE, size=10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_exercice_guide(doc, titre, notion, enonce_lines, methode_steps, verification, conclusion):
    """
    Exercice guidé complet — fond gris très clair, titre bleu.
    notion         : str — notion ciblée
    enonce_lines   : liste de str — énoncé
    methode_steps  : liste de str — étapes (sans numéro, la fonction les numérote)
    verification   : str
    conclusion     : str ou liste de str
    """
    t, cell = _full_width_1cell(doc)
    _set_cell_bg(cell, "F5F7FA")
    _set_cell_borders(cell, BLEU, sz=6, sides=("top","left","bottom","right"))
    _set_cell_margins(cell, top=100, bottom=100, left=160, right=160)
    p = cell.paragraphs[0]; p.text = ""
    r(p, "◉  " + titre, bold=True, color=BLEU, size=11)
    pp = cell.add_paragraph(); pp.paragraph_format.space_after = Pt(2)
    r(pp, "Notion ciblée : ", bold=True, color=OR, size=10.5)
    r(pp, notion, italic=True, color=TEXTE, size=10.5)

    # Énoncé
    pp2 = cell.add_paragraph(); pp2.paragraph_format.space_before = Pt(4)
    r(pp2, "Énoncé :", bold=True, color=TEXTE, size=10.5)
    for ln in enonce_lines:
        pp3 = cell.add_paragraph(); pp3.paragraph_format.space_after = Pt(1)
        r(pp3, ln, color=TEXTE, size=10.5)

    # Méthode
    pp4 = cell.add_paragraph(); pp4.paragraph_format.space_before = Pt(4)
    r(pp4, "MÉTHODE", bold=True, color=BLEU, size=10.5)
    for i, step in enumerate(methode_steps, 1):
        pp5 = cell.add_paragraph(); pp5.paragraph_format.space_after = Pt(1)
        r(pp5, f"  Étape {i} · ", bold=True, color=BLEU, size=10.5)
        r(pp5, step, color=TEXTE, size=10.5)

    # Vérification
    ppv = cell.add_paragraph(); ppv.paragraph_format.space_before = Pt(4)
    r(ppv, "✔ Vérification : ", bold=True, color=VERT, size=10.5)
    r(ppv, verification, color=TEXTE, size=10.5)

    # Conclusion
    ppc = cell.add_paragraph(); ppc.paragraph_format.space_after = Pt(2)
    r(ppc, "✔ Résultats : ", bold=True, color=VERT, size=10.5)
    if isinstance(conclusion, str):
        r(ppc, conclusion, color=TEXTE, size=10.5)
    else:
        for ln in conclusion:
            ppc2 = cell.add_paragraph(); ppc2.paragraph_format.space_after = Pt(1)
            r(ppc2, "     " + ln, color=TEXTE, size=10.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_exercice(doc, titre, lignes, guide=None):
    """Exercice d'entraînement avec cadre léger + guide optionnel."""
    t, cell = _full_width_1cell(doc)
    _set_cell_bg(cell, "FAFAFA")
    _set_cell_borders(cell, "CCCCCC", sz=4, sides=("top","left","bottom","right"))
    _set_cell_margins(cell, top=100, bottom=100, left=160, right=160)
    p = cell.paragraphs[0]; p.text = ""
    r(p, titre, bold=True, color=BLEU, size=11)
    for ln in lignes:
        pp = cell.add_paragraph(); pp.paragraph_format.space_after = Pt(2)
        if isinstance(ln, str):
            r(pp, ln, color=TEXTE, size=10.5)
        else:
            for item in ln:
                if isinstance(item, str):
                    r(pp, item, color=TEXTE, size=10.5)
                else:
                    txt, style = item
                    r(pp, txt, bold=("b" in style), italic=("i" in style),
                      underline=("u" in style), color=TEXTE, size=10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    if guide:
        add_encadre_bleu(doc, "Guide de réussite", [guide])


def add_devoir_maison(doc, titre, duree, enonce_lines, guide_steps):
    """
    Devoir maison — encadré or/amber avec guide pas à pas.
    titre        : str
    duree        : str (ex: "30 minutes")
    enonce_lines : liste de str — contexte + consignes
    guide_steps  : liste de str — étapes sans numéro
    """
    t, cell = _full_width_1cell(doc)
    _set_cell_bg(cell, "FFF9EE")
    _set_cell_borders(cell, OR, sz=12, sides=("left",))
    _set_cell_borders(cell, JAUNE_BG, sz=4, sides=("top","bottom","right"))
    _set_cell_margins(cell, top=100, bottom=100, left=180, right=160)
    p = cell.paragraphs[0]; p.text = ""
    r(p, "◈  DEVOIR MAISON — " + titre, bold=True, color=OR, size=11)
    pp = cell.add_paragraph()
    r(pp, f"Durée conseillée : {duree}", italic=True, color=GRIS, size=10)

    # Énoncé
    for ln in enonce_lines:
        pp2 = cell.add_paragraph(); pp2.paragraph_format.space_after = Pt(2)
        if isinstance(ln, tuple) and ln[0] == "TITRE":
            r(pp2, ln[1], bold=True, color=OR, size=10.5)
        else:
            r(pp2, ln, color=TEXTE, size=10.5)

    # Guide
    pp3 = cell.add_paragraph(); pp3.paragraph_format.space_before = Pt(4)
    r(pp3, "GUIDE PAS À PAS (sans corrigé)", bold=True, color=OR, size=10.5)
    for i, step in enumerate(guide_steps, 1):
        pp4 = cell.add_paragraph(); pp4.paragraph_format.space_after = Pt(1)
        r(pp4, f"  Étape {i} · ", bold=True, color=OR, size=10.5)
        r(pp4, step, color=TEXTE, size=10.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_corrige(doc, titre, lignes):
    """Encadré de corrigé — fond vert très clair."""
    t, cell = _full_width_1cell(doc)
    _set_cell_bg(cell, "F8FBF7")
    _set_cell_borders(cell, VERT, sz=4, sides=("top","left","bottom","right"))
    _set_cell_margins(cell, top=100, bottom=100, left=160, right=160)
    p = cell.paragraphs[0]; p.text = ""
    r(p, "✔  " + titre, bold=True, color=VERT, size=11)
    for ln in lignes:
        pp = cell.add_paragraph(); pp.paragraph_format.space_after = Pt(2)
        if isinstance(ln, str):
            r(pp, ln, color=TEXTE, size=10.5)
        else:
            for item in ln:
                if isinstance(item, str):
                    r(pp, item, color=TEXTE, size=10.5)
                else:
                    txt, style = item
                    r(pp, txt, bold=("b" in style), italic=("i" in style),
                      underline=("u" in style), color=TEXTE, size=10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


# -------------------------------------------------------
# Tableaux
# -------------------------------------------------------
def add_table_2col(doc, header, rows_data, w1=Cm(6), w2=Cm(10.5)):
    """Tableau 2 colonnes avec en-tête bleu foncé."""
    t = doc.add_table(rows=1, cols=2)
    t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    t.columns[0].width = w1; t.columns[1].width = w2
    hdr = t.rows[0].cells
    for i, h in enumerate(header):
        hdr[i].paragraphs[0].text = ""
        r(hdr[i].paragraphs[0], h, bold=True, color="FFFFFF", size=10)
        _set_cell_bg(hdr[i], BLEU_FONCE)
        _set_cell_margins(hdr[i], top=80, bottom=80, left=120, right=120)
    for row_vals in rows_data:
        row = t.add_row().cells
        for i, val in enumerate(row_vals):
            row[i].paragraphs[0].text = ""
            if isinstance(val, str):
                r(row[i].paragraphs[0], val, color=TEXTE, size=10)
            else:
                for item in val:
                    if isinstance(item, str):
                        r(row[i].paragraphs[0], item, color=TEXTE, size=10)
                    else:
                        txt, style = item
                        r(row[i].paragraphs[0], txt,
                          bold=("b" in style), italic=("i" in style),
                          underline=("u" in style), color=TEXTE, size=10)
            _set_cell_margins(row[i], top=60, bottom=60, left=120, right=120)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


# -------------------------------------------------------
# Schéma / Mindmap matplotlib
# -------------------------------------------------------
def make_mindmap(path, center, branches):
    fig, ax = plt.subplots(figsize=(8, 5.2)); ax.axis("off")
    ax.set_xlim(0,10); ax.set_ylim(0,10)
    palette = ["#184082","#C99225","#2A7E5B","#C0541A","#6B3FA0"]
    ax.add_patch(FancyBboxPatch((3.5,4.2), 3.0, 1.6,
                 boxstyle="round,pad=0.1", fc="#184082", ec="none"))
    ax.text(5, 5, center, ha="center", va="center", color="white",
            fontsize=11, fontweight="bold")
    n = len(branches)
    angles = np.linspace(90, 90+360, n, endpoint=False)
    for i, (label, subs) in enumerate(branches):
        col = palette[i % len(palette)]
        a = np.deg2rad(angles[i])
        bx, by = 5+3.3*np.cos(a), 5+2.9*np.sin(a)
        ax.plot([5+1.5*np.cos(a), bx],[5+0.8*np.sin(a), by], color=col, lw=2)
        ax.add_patch(FancyBboxPatch((bx-1.1,by-0.38), 2.2, 0.76,
                     boxstyle="round,pad=0.06", fc="white", ec=col, lw=2))
        ax.text(bx, by, label, ha="center", va="center",
                color=col, fontsize=9.5, fontweight="bold")
        for j, s in enumerate(subs):
            off = (j-(len(subs)-1)/2)*0.55
            sx = bx+1.7*np.cos(a); sy = by+off
            ax.plot([bx+1.1*np.cos(a), sx-0.05],[by, sy], color=col, lw=1, alpha=0.5)
            ax.text(sx, sy, "• "+s, ha="left", va="center",
                    color="#333333", fontsize=8)
    plt.tight_layout()
    plt.savefig(path, dpi=200, bbox_inches="tight", facecolor="white")
    plt.close()


def make_schema_present_continuous(path):
    """Schéma visuel de formation du Present Continuous."""
    fig, ax = plt.subplots(figsize=(9, 5)); ax.axis("off")
    ax.set_xlim(0, 10); ax.set_ylim(0, 6)

    # Titre
    ax.text(5, 5.5, "PRESENT CONTINUOUS — Formation", ha="center", va="center",
            fontsize=13, fontweight="bold", color="#184082")

    # Rectangle central (formule)
    ax.add_patch(FancyBboxPatch((2.5, 4.0), 5.0, 0.8,
                 boxstyle="round,pad=0.1", fc="#D9E2F3", ec="#184082", lw=2))
    ax.text(5, 4.4, "Subject  +  AM / IS / ARE  +  Verb + -ING", ha="center",
            va="center", fontsize=11, fontweight="bold", color="#183E82")

    # Colonne gauche — sujets
    sujets = [("I", "am", 3.0), ("He / She / It", "is", 2.2), ("You / We / They", "are", 1.4)]
    colors = ["#2A7E5B", "#184082", "#C99225"]
    for (subj, aux, y), col in zip(sujets, colors):
        ax.add_patch(FancyBboxPatch((0.2, y-0.28), 2.6, 0.56,
                     boxstyle="round,pad=0.05", fc="white", ec=col, lw=1.5))
        ax.text(1.5, y, f"{subj}  →  {aux}", ha="center", va="center",
                fontsize=10, color=col, fontweight="bold")

    # Flèches vers formule
    for y in [3.0, 2.2, 1.4]:
        ax.annotate("", xy=(2.5, 3.8), xytext=(2.8, y),
                    arrowprops=dict(arrowstyle="->", color="#555555", lw=1))

    # Colonne droite — exemples
    exemples = [
        ("I am walking.", "Je marche.", 3.0),
        ("She is wearing.", "Elle porte.", 2.2),
        ("They are clapping.", "Ils applaudissent.", 1.4),
    ]
    for (en, fr, y) in exemples:
        ax.text(5.5, y+0.12, en, ha="left", va="center", fontsize=9.5,
                color="#232323", fontweight="bold")
        ax.text(5.5, y-0.15, fr, ha="left", va="center", fontsize=8.5,
                color="#555555", style="italic")

    # Encadré piège
    ax.add_patch(FancyBboxPatch((0.1, 0.2), 4.0, 0.8,
                 boxstyle="round,pad=0.05", fc="#FDECEA", ec="#A62D2D", lw=1.5))
    ax.text(2.1, 0.6, "✗  I am wear a dress.", ha="center", va="center",
            fontsize=9, color="#A62D2D", fontweight="bold")
    ax.add_patch(FancyBboxPatch((4.8, 0.2), 5.0, 0.8,
                 boxstyle="round,pad=0.05", fc="#E2F0D9", ec="#2A7E5B", lw=1.5))
    ax.text(7.3, 0.6, "✓  I am wearing a dress.", ha="center", va="center",
            fontsize=9, color="#2A7E5B", fontweight="bold")

    ax.text(5, -0.1, "Action qui se passe EN CE MOMENT", ha="center", va="center",
            fontsize=9, color="#555555", style="italic")

    plt.tight_layout()
    plt.savefig(path, dpi=200, bbox_inches="tight", facecolor="white")
    plt.close()


def add_figure(doc, path, caption, width_cm=14):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path, width=Cm(width_cm))
    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r(cap, caption, italic=True, color=GRIS, size=9)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


# -------------------------------------------------------
# En-tête / Pied de page
# -------------------------------------------------------
def _add_page_field(p):
    run = p.add_run(); run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor.from_string(GRIS)
    fld1 = OxmlElement("w:fldChar"); fld1.set(qn("w:fldCharType"),"begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"),"preserve"); instr.text = "PAGE"
    fld2 = OxmlElement("w:fldChar"); fld2.set(qn("w:fldCharType"),"end")
    run._r.append(fld1); run._r.append(instr); run._r.append(fld2)

def setup_header_footer(doc, niveau, serie, matiere_label, lecon, prof):
    from docx.enum.text import WD_TAB_ALIGNMENT
    sec = doc.sections[0]
    h = sec.header; h.is_linked_to_previous = False
    hp = h.paragraphs[0]; hp.text = ""
    hp.paragraph_format.tab_stops.add_tab_stop(Cm(16.5), WD_TAB_ALIGNMENT.RIGHT)
    left = f"LPA · {niveau}" + (f" · {serie}" if serie else "")
    r(hp, left, size=9, color=GRIS)
    r(hp, f"\t{matiere_label} · Leçon {lecon}", size=9, color=BLEU, bold=True)
    _para_border(hp, BLEU, side="bottom", sz=12)
    f = sec.footer; f.is_linked_to_previous = False
    fp = f.paragraphs[0]; fp.text = ""
    fp.paragraph_format.tab_stops.add_tab_stop(Cm(8.25), WD_TAB_ALIGNMENT.CENTER)
    fp.paragraph_format.tab_stops.add_tab_stop(Cm(16.5), WD_TAB_ALIGNMENT.RIGHT)
    _para_border(fp, BLEU, side="top", sz=6)
    r(fp, f"Professeur {prof}", size=8.5, color=GRIS)
    r(fp, "\tDPFC/MENET-FP CI · 2025-2026", size=8.5, color=GRIS)
    r(fp, "\tPage ", size=8.5, color=GRIS)
    _add_page_field(fp)


# -------------------------------------------------------
# Page de garde
# -------------------------------------------------------
def add_cover(doc, titre_lecon, sous_titre, niveau, serie,
              matiere_label, num_lecon, prof, coef):
    for _ in range(3): doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r(p, "LPA", bold=True, color=BLEU, size=48)
    p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r(p2, "Lycée Personnel Autonome", italic=True, color=GRIS, size=12)
    doc.add_paragraph()
    sep = doc.add_paragraph(); sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_border(sep, BLEU, side="bottom", sz=16)
    doc.add_paragraph()
    pm = doc.add_paragraph(); pm.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ligne_meta = (f"{matiere_label}  ·  Leçon {num_lecon}  ·  {niveau}"
                  + (f"  ·  {serie}" if serie else ""))
    r(pm, ligne_meta, bold=True, color=BLEU, size=13)
    pt = doc.add_paragraph(); pt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pt.paragraph_format.space_before = Pt(12)
    r(pt, titre_lecon, bold=True, color=OR, size=26)
    if sous_titre:
        ps = doc.add_paragraph(); ps.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r(ps, sous_titre, italic=True, color=GRIS, size=13)
    doc.add_paragraph()
    pi = doc.add_paragraph(); pi.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r(pi, f"Professeur {prof}", color=TEXTE, size=11)
    r(pi, f"  ·  Coefficient {coef}", color=GRIS, size=11)
    for _ in range(5): doc.add_paragraph()
    pd = doc.add_paragraph(); pd.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r(pd, "DPFC / MENET-FP Côte d'Ivoire  ·  2025-2026", color=GRIS, size=10)
    doc.add_page_break()


# ============================================================
# SECTION 2 — CONTENU — LEÇON 12 : FASHION SHOW
# ============================================================

TITRE       = "Fashion show — Le défilé de mode"
SOUS_TITRE  = "Compétence C4 · Écrit élaboré · 3ème"
NIVEAU      = "3ème"
SERIE       = None
MATIERE     = "Anglais LV1"
LECON       = 12
PROF        = "Evelyne ASSI"
COEF        = 2

doc = Document()
st = doc.styles["Normal"]; st.font.name = "Arial"; st.font.size = Pt(11)
st.paragraph_format.line_spacing = 1.15
st.paragraph_format.space_after  = Pt(6)
sec = doc.sections[0]
sec.top_margin = Cm(2.5); sec.bottom_margin = Cm(2.5)
sec.left_margin = Cm(2);  sec.right_margin  = Cm(2)

setup_header_footer(doc, NIVEAU, SERIE, MATIERE, LECON, PROF)
add_cover(doc, TITRE, SOUS_TITRE, NIVEAU, SERIE, MATIERE, LECON, PROF, COEF)

# ============================================================
# PARTIE 1 — AVANT LA LEÇON
# ============================================================
add_part_banner(doc, "Section 1 — Avant la leçon")

# --- Capsule de reprise ---
add_capsule_reprise(doc,
    points_5=[
        "On distingue modern clothes (vêtements modernes) et traditional clothes (vêtements traditionnels).",
        "Vocabulaire clé : dress (robe), shirt (chemise), trousers (pantalon), kente, boubou.",
        "Le Present Simple décrit les habitudes vestimentaires : \"She wears a boubou every Friday.\"",
        "L'adjectif se place AVANT le nom en anglais : \"a beautiful dress\" — jamais \"a dress beautiful\".",
        "Verbes utiles : wear (porter), dress (s'habiller), choose (choisir), design (concevoir).",
    ],
    questions_3=[
        ("Comment dit-on \"tissu traditionnel\" en anglais ?", "traditional fabric"),
        ("Traduis : \"He is wearing a white shirt.\"", "Il porte une chemise blanche."),
        ("Place l'adjectif correctement : shoes / leather / brown", "brown leather shoes"),
    ],
    lien="Aujourd'hui on ne parle plus seulement des vêtements : on décrit un ÉVÉNEMENT — le défilé de mode. "
         "Tu vas apprendre à rédiger une description en anglais et utiliser le Present Continuous pour décrire ce qui se passe en ce moment."
)

# --- Ancrage ivoirien ---
add_ancrage_ivoirien(doc, [
    "Imagine que tu es à Abidjan, au Palais de la Culture de Treichville. Ce soir a lieu le Grand Défilé Wax & Fashion, organisé par le styliste ivoirien Adama TRAORÉ.",
    "Des mannequins défilent sur le podium vêtus de tenues en wax coloré, en bogolan et en bazin riche. Des centaines de spectateurs applaudissent. Des journalistes de RTI prennent des notes.",
    "Toi, tu dois écrire un article en anglais pour le journal de ton école sur cet événement.",
    "Lien DPFC : Entrepreneuriat et valorisation du patrimoine culturel ivoirien à travers la mode.",
])

# --- Harkness ---
add_harkness(doc,
    question="Can fashion show who you are and where you come from? (La mode peut-elle montrer qui tu es et d'où tu viens ?)",
    guide="Pense aux vêtements traditionnels de ta région (wax, kente, boubou…). / Pense aux vêtements que tu portes à l'école vs ceux du week-end. / Est-ce qu'une tenue peut raconter une histoire ?",
    corrige_txt="Yes, fashion can show who you are and where you come from. Traditional clothes like the wax or the boubou show African identity and pride. In Côte d'Ivoire, many people mix both styles — wearing a wax top with jeans. (Oui, la mode peut montrer qui tu es. Les vêtements traditionnels montrent l'identité africaine. Beaucoup de personnes mélangent les deux styles.)"
)
add_harkness(doc,
    question="How do you describe what is happening right now in English? (Comment décrit-on ce qui se passe en ce moment en anglais ?)",
    guide="Quelle est la différence entre \"she walks\" et \"she is walking\" ? / Quand utilise-t-on le Present Continuous ? / Comment forme-t-on le Present Continuous ?",
    corrige_txt="We use the Present Continuous to describe actions happening right now. Form : Subject + am/is/are + verb + -ing. Example : \"The model is walking on the runway.\" We use the Present Simple for habits : \"She walks every day.\" (On utilise le Present Continu pour les actions en cours. Forme : Sujet + am/is/are + verbe-ing. Exemple : Le mannequin marche sur le podium — en ce moment.)"
)
add_harkness(doc,
    question="What words do you need to describe a fashion show? (Quels mots faut-il pour décrire un défilé de mode ?)",
    guide="Qui participe à un défilé de mode ? / Qu'est-ce qu'on voit, qu'est-ce qu'on entend ? / Comment décrire une tenue de façon précise ?",
    corrige_txt="People : model (mannequin), designer (styliste), audience (public), photographer (photographe). Place : runway (podium), backstage (coulisses), spotlight (projecteur). Adjectives : elegant (élégant), colourful (coloré), fitted (ajusté), loose (ample), embroidered (brodé), patterned (à motifs)."
)

# ============================================================
# PARTIE 2 — LA LEÇON
# ============================================================
add_part_banner(doc, "Section 2 — La leçon")

# --- Phase 1 ---
add_phase_title(doc, "Phase 1 · Présentation / Prérequis")
add_body(doc,
    "Professeur Evelyne ASSI : ",
    ("\"Aujourd'hui on ne se contente pas de nommer les vêtements — on raconte ce qui se passe. Parle même si c'est imparfait — le silence n'apprend rien.\"", "i")
)
add_subsection(doc, "Objectifs de la leçon")
add_bullet(doc, "Décrire un événement de mode en anglais (à l'écrit)")
add_bullet(doc, "Utiliser le Present Continuous correctement")
add_bullet(doc, "Maîtriser le vocabulaire du défilé de mode")
add_bullet(doc, "Rédiger un paragraphe descriptif structuré")
add_subsection(doc, "Prérequis nécessaires")
add_bullet(doc, "Vocabulaire des vêtements (Leçon 11)")
add_bullet(doc, "Les adjectifs de couleur, de taille, de matière")
add_bullet(doc, "La structure Subject + Verb + Complement")
add_bullet(doc, "Le verbe \"to be\" au présent : am / is / are")

# --- Phase 2 ---
add_phase_title(doc, "Phase 2 · Découverte guidée")
add_body(doc, "Lis ce texte — il se passe en ce moment à Abidjan :")
add_encadre_bleu(doc, "Texte de découverte — Grand Wax Fashion Show", [
    "It is Friday evening in Abidjan. The Grand Wax Fashion Show is taking place at the Palais de la Culture de Treichville.",
    "The audience is clapping and cheering. On the runway, a model named Mariama is walking confidently.",
    "She is wearing a long, colourful wax dress with golden embroidery. Another model, Kouamé, is showing a modern boubou in deep blue bogolan fabric.",
    "The designer, Adama Traoré, is smiling backstage. Photographers are taking pictures. Everyone is enjoying the show.",
    "— Traduction : C'est vendredi soir à Abidjan. Le Grand Défilé Wax Fashion se déroule au Palais de la Culture. Le public applaudit. Sur le podium, Mariama marche avec assurance. Elle porte une longue robe en wax colorée avec broderie dorée. Kouamé présente un boubou bogolan bleu profond. Le styliste sourit dans les coulisses. Tout le monde apprécie le spectacle.",
])
add_retenir(doc, "À RETENIR — Present Continuous", [
    "Le PRESENT CONTINUOUS décrit une action qui se passe EN CE MOMENT.",
    "FORME : Subject + am / is / are + verb + -ING",
    "· I am watching the show.  (Je regarde le spectacle.)",
    "· She is wearing a wax dress.  (Elle porte une robe en wax.)",
    "· They are taking pictures.  (Ils prennent des photos.)",
    "CONTRAIRE : le Present Simple décrit une HABITUDE.",
    "· She wears wax every Friday.  (Elle porte du wax chaque vendredi.)",
])
add_subsection(doc, "Formation du -ING — Règle complète")
add_table_2col(doc,
    ["Règle", "Exemples"],
    [
        ("Verbe en -E muet → retire le E, ajoute -ING", "make → making  /  write → writing  /  smile → smiling"),
        ("Verbe court CVC (consonne-voyelle-consonne) → double la consonne finale, ajoute -ING", "run → running  /  sit → sitting  /  clap → clapping"),
        ("Tous les autres verbes → ajoute simplement -ING", "walk → walking  /  wear → wearing  /  show → showing"),
    ],
    w1=Cm(7), w2=Cm(9.5)
)
add_attention(doc, "Faux amis et pièges", [
    "\"He is wearing\" = Il porte (en ce moment) — PAS \"il est portant\". Ne traduis jamais mot à mot.",
    "\"Fashion\" se prononce /ˈfæʃ.ən/ — FA-cheune — PAS \"fachione\".",
    "\"Designer\" = styliste/créateur de mode. Prononciation : /dɪˈzaɪ.nər/ — di-ZAI-neur.",
    "\"Runway\" = podium/passerelle (les deux termes sont acceptés en français).",
    "\"Outfit\" = une tenue COMPLÈTE — pas seulement un vêtement seul.",
])

# --- Phase 3 — Schémas ---
add_phase_title(doc, "Phase 3 · Schémas textuels")
make_schema_present_continuous("/home/claude/schema_present_continuous.png")
add_figure(doc, "/home/claude/schema_present_continuous.png",
           "Schéma 1 — Formation du Present Continuous")
make_mindmap("/home/claude/mindmap_fashion.png",
    "FASHION\nSHOW",
    [
        ("PEOPLE",     ["model (mannequin)", "designer (styliste)", "audience (public)", "photographer"]),
        ("PLACE",      ["runway (podium)", "backstage (coulisses)", "spotlight (projecteur)", "stage (scène)"]),
        ("CLOTHES",    ["outfit (tenue)", "collection", "embroidery (broderie)", "fabric (tissu)"]),
        ("ADJECTIVES", ["elegant (élégant)", "colourful (coloré)", "fitted (ajusté)", "patterned (à motifs)"]),
    ]
)
add_figure(doc, "/home/claude/mindmap_fashion.png",
           "Schéma 2 — Carte mentale : Vocabulary of a Fashion Show")

# --- Phase 4 — Définitions DPFC ---
add_phase_title(doc, "Phase 4 · Définitions DPFC")
add_encadre_bleu(doc, "DÉFINITION DPFC — Fashion show", [
    "A fashion show is an event where designers present their new collection of clothes through models walking on a runway.",
    "(Un défilé de mode est un événement où les stylistes présentent leur nouvelle collection de vêtements à travers des mannequins défilant sur un podium.)",
])
add_encadre_bleu(doc, "DÉFINITION DPFC — Present Continuous", [
    "The Present Continuous is a verb tense used to describe an action happening at the moment of speaking.",
    "It is formed with : Subject + am/is/are + verb + -ing.",
    "(Le Present Continu est un temps verbal pour décrire une action qui se passe au moment où l'on parle. Forme : Sujet + am/is/are + verbe + -ing.)",
])
add_encadre_bleu(doc, "DÉFINITION DPFC — Model / Designer / Outfit", [
    "Model : a person who wears clothes during a fashion show to present a designer's collection. (Mannequin : personne qui présente la collection d'un styliste.)",
    "Designer : a person who creates and makes clothes, bags, shoes or accessories. (Styliste : créateur de vêtements, sacs, chaussures ou accessoires.)",
    "Outfit : a set of clothes worn together on a specific occasion. (Tenue : ensemble de vêtements portés pour une occasion précise.)",
])

# --- Phase 5 — Trace écrite + Analogies ---
add_phase_title(doc, "Phase 5 · Trace écrite + Analogies CI")
add_trace_ecrite(doc, [
    "=== LEÇON 12 — FASHION SHOW (Le défilé de mode)",
    "Compétence C4 · Écrit élaboré · Anglais LV1 · 3ème",
    "",
    "=== I. VOCABULAIRE DU DÉFILÉ DE MODE",
    "Personnes : model = mannequin  |  designer = styliste  |  audience = public  |  photographer = photographe",
    "Lieux et objets : runway = podium/passerelle  |  backstage = coulisses  |  outfit = tenue  |  fabric = tissu",
    "Adjectifs : elegant = élégant  |  colourful = coloré  |  fitted = ajusté  |  loose = ample  |  embroidered = brodé  |  patterned = à motifs",
    "",
    "=== II. LE PRESENT CONTINUOUS",
    "Utilisation : action qui se passe EN CE MOMENT.",
    "Formation : Subject + am / is / are + verb + -ING",
    "I am walking.        → Je marche.         (en ce moment)",
    "She is wearing.      → Elle porte.         (en ce moment)",
    "They are clapping.   → Ils applaudissent.  (en ce moment)",
    "Règles du -ING : verbe en -E muet → retire le E (make→making) | verbe court CVC → double la consonne (run→running) | autres → ajoute -ING (walk→walking)",
    "",
    "=== III. STRUCTURE POUR DÉCRIRE UNE TENUE",
    "[Name] is wearing + a/an + [adj. couleur] + [adj. style] + [nom du vêtement] + with + [détail].",
    "Exemple : \"Mariama is wearing a long, colourful wax dress with golden embroidery.\"",
    "(Mariama porte une longue robe en wax colorée avec une broderie dorée.)",
])
add_analogie(doc, [
    "Analogie 1 — Le Present Continuous, c'est comme un commentateur sportif sur Radio Côte d'Ivoire pendant un match des Éléphants : il décrit CE QUI SE PASSE EN CE MOMENT, pas ce qui se passe d'habitude.",
    "\"Gradel is running!\" = Gradel court ! (maintenant, sur le terrain) — PAS \"Gradel runs\" qui voudrait dire qu'il court tous les jours.",
    "",
    "Analogie 2 — Décrire une tenue à un défilé, c'est comme décrire l'habit d'un chef traditionnel à une cérémonie de chefferie à Yamoussoukro : on donne la couleur, le tissu, les ornements — dans cet ordre précis.",
])

# --- Phase 5b — Synthèse ---
add_phase_title(doc, "Phase 5b · Synthèse")
add_synthese(doc,
    points=[
        "Un fashion show est un événement où des mannequins présentent des collections de vêtements.",
        "On utilise le Present Continuous pour décrire une action en cours : Subject + am/is/are + verb-ing.",
        "Pour décrire une tenue, l'ordre est : couleur + style + nom du vêtement + détail.",
        "Les adjectifs se placent TOUJOURS avant le nom en anglais.",
        "Un texte descriptif d'événement suit l'ordre : Qui → Quoi → Comment.",
    ],
    mots_cles=[
        ("Fashion show", "défilé de mode — événement de présentation de vêtements"),
        ("Model",        "mannequin — personne qui présente les tenues"),
        ("Designer",     "styliste — créateur de vêtements"),
        ("Outfit",       "tenue — ensemble de vêtements"),
        ("Runway",       "podium — passerelle où défilent les mannequins"),
    ],
    erreurs=[
        ("She is wear a dress.",          "She is WEARing a dress. (toujours le -ING !)"),
        ("a dress colourful",             "a COLOURFUL dress (adjectif avant le nom)"),
        ("She wearing a dress.",          "She IS wearing a dress. (am/is/are obligatoire)"),
    ]
)

# --- Phase 6 — Exercices guidés ---
add_phase_title(doc, "Phase 6 · Exercices guidés")
add_exercice_guide(doc,
    titre="Exercice guidé 1 — Conjugaison au Present Continuous",
    notion="Formation du Present Continuous",
    enonce_lines=[
        "Mets les verbes entre parenthèses au Present Continuous.",
        "1. The model ________ (walk) on the runway.",
        "2. The designer ________ (smile) backstage.",
        "3. The photographers ________ (take) pictures.",
        "4. Aïcha ________ (wear) a beautiful wax dress.",
        "5. We ________ (watch) the fashion show.",
    ],
    methode_steps=[
        "Identifie le sujet de chaque phrase → \"The model\" = he/she → is / \"The photographers\" = they → are / \"We\" = we → are",
        "Applique la règle du -ING : walk→walking (ordinaire) / smile→smiling (-E muet) / take→taking (-E muet) / wear→wearing (ordinaire) / watch→watching (ordinaire)",
        "Assemble : Subject + is/are + verb-ING",
    ],
    verification="chaque phrase doit contenir am/is/are + verbe-ING",
    conclusion=[
        "1. The model is walking on the runway.  (Le mannequin marche sur le podium.)",
        "2. The designer is smiling backstage.  (Le styliste sourit dans les coulisses.)",
        "3. The photographers are taking pictures.  (Les photographes prennent des photos.)",
        "4. Aïcha is wearing a beautiful wax dress.  (Aïcha porte une belle robe en wax.)",
        "5. We are watching the fashion show.  (Nous regardons le défilé de mode.)",
    ]
)
add_exercice_guide(doc,
    titre="Exercice guidé 2 — Rédiger une description de tenue",
    notion="Structure de description d'une tenue + vocabulaire du défilé",
    enonce_lines=[
        "Regarde ces informations sur le mannequin Kouamé et écris une description complète de 3 phrases.",
        "· Prénom : Kouamé  |  Il marche sur le podium (en ce moment)",
        "· Tenue : boubou / bleu foncé / tissu bogolan / broderies argentées",
        "· Le public applaudit (en ce moment)",
    ],
    methode_steps=[
        "Phrase 1 — qui + action : Kouamé + is + walking + on the runway",
        "Phrase 2 — décrire la tenue : ordre → couleur + tissu + vêtement + détail → \"He is wearing a dark blue bogolan boubou with silver embroidery.\"",
        "Phrase 3 — réaction du public : \"The audience\" → are (pluriel collectif) + clap → clapping (CVC : double P)",
    ],
    verification="3 phrases au Present Continuous · vocabulaire du défilé présent · adjectifs avant les noms",
    conclusion=[
        "\"Kouamé is walking on the runway. He is wearing a dark blue bogolan boubou with silver embroidery. The audience is clapping.\"",
    ]
)

# ============================================================
# PARTIE 3 — APRÈS LA LEÇON
# ============================================================
add_part_banner(doc, "Section 3 — Après la leçon")

add_exercice(doc,
    "◎ Exercice 1 — Vocabulaire du défilé de mode · Notions : fashion show vocabulary",
    [
        "Relie chaque mot anglais à sa traduction française.",
        "Colonne A — 1. runway  |  2. designer  |  3. audience  |  4. outfit  |  5. backstage",
        "Colonne B — a. styliste  |  b. coulisses  |  c. tenue  |  d. podium/passerelle  |  e. public/spectateurs",
        "Écris tes réponses sous la forme : 1→d, 2→a, etc.",
    ],
    guide="Étape 1 : relis la trace écrite — vocabulaire. Étape 2 : cherche la traduction dans la Colonne B. Étape 3 : note 1→?, 2→?, etc."
)
add_exercice(doc,
    "◎ Exercice 2 — Formation du -ING · Notions : règles du -ING / conjugaison",
    [
        "Transforme ces verbes à l'infinitif en forme -ING. Indique la règle appliquée.",
        "1. dance  →  _______   (règle : ?)",
        "2. sit    →  _______   (règle : ?)",
        "3. show   →  _______   (règle : ?)",
        "4. write  →  _______   (règle : ?)",
        "5. run    →  _______   (règle : ?)",
    ],
    guide="Étape 1 : identifie la terminaison de chaque verbe (-E muet ? CVC ? ordinaire ?). Étape 2 : applique la règle. Étape 3 : écris la forme -ING."
)
add_exercice(doc,
    "◎ Exercice 3 — Construire des phrases · Notions : Subject + am/is/are + verb-ING",
    [
        "Construis des phrases complètes au Present Continuous avec les éléments donnés.",
        "1. Fatou / wear / a long red dress",
        "2. The photographer / take / pictures",
        "3. I / look / at the runway",
        "4. The models / walk / on the stage",
        "5. The designer / show / his new collection",
    ],
    guide="Étape 1 : identifie le sujet → choisis am/is/are. Étape 2 : mets le verbe à la forme -ING. Étape 3 : ajoute le complément."
)
add_exercice(doc,
    "◎ Exercice 4 — Décrire une tenue · Notions : ordre des adjectifs + vocabulaire",
    [
        "Lis ces informations et écris une phrase de description pour chaque mannequin. Utilise le Present Continuous.",
        "Awa : robe / wax / verte / longue / avec ceinture dorée",
        "Issiaka : boubou / blanc / en bazin / avec broderies bleues",
        "Nadia : jupe / courte / noire / avec motifs géométriques",
    ],
    guide="Étape 1 : structure → [Name] is wearing a/an [adj.] [adj.] [vêtement] with [détail]. Étape 2 : adjectifs AVANT le nom. Étape 3 : ajoute le détail avec \"with\"."
)
add_exercice(doc,
    "◎ Exercice 5 — Rédiger un paragraphe descriptif · Notions : cohérence textuelle + Present Continuous + vocabulaire complet",
    [
        "Tu assistes au Grand Défilé Wax Fashion au Palais de la Culture d'Abidjan.",
        "Rédige un paragraphe de 5 phrases en anglais pour décrire ce que tu vois en ce moment.",
        "Tu dois inclure : au moins 3 verbes au Present Continuous · au moins 2 descriptions de tenues",
        "· au moins 3 mots de vocabulaire du défilé (runway, designer, audience, outfit, backstage…)",
        "· le nom d'un lieu ivoirien et un prénom ivoirien",
    ],
    guide="Étape 1 : situe l'action (où ? quand ?). Étape 2 : décris ce qui se passe sur le podium. Étape 3 : décris les tenues de 2 mannequins. Étape 4 : décris la réaction du public. Étape 5 : vérifie tes formes au Present Continuous."
)

# --- Devoir Maison ---
add_devoir_maison(doc,
    titre="A Fashion Show in Abidjan",
    duree="30 minutes",
    enonce_lines=[
        ("TITRE", "CONTEXTE :"),
        "Tu es journaliste pour le journal de ton lycée à Abidjan. Tu viens d'assister au premier défilé de mode organisé par des élèves de Terminale. Tu dois écrire un article court en anglais.",
        ("TITRE", "CONSIGNE :"),
        "Rédige un article en anglais de 80 à 100 mots contenant :",
        "1. Un titre en anglais.",
        "2. Une introduction : où et quand l'événement a-t-il eu lieu ?",
        "3. Un développement : décris 2 mannequins et leurs tenues (Present Continuous).",
        "4. Une conclusion : comment a réagi le public ?",
        ("TITRE", "CONTRAINTES OBLIGATOIRES :"),
        "· Minimum 4 phrases au Present Continuous",
        "· Minimum 4 mots de vocabulaire du défilé",
        "· Au moins 1 prénom ivoirien et 1 lieu d'Abidjan",
        "· Les adjectifs AVANT les noms",
    ],
    guide_steps=[
        "Trouve un titre accrocheur. Ex : \"A Wonderful Fashion Show at [Nom de ton lycée]\"",
        "Rédige l'introduction (1-2 phrases) : Quand ? Où ? Qui organise ?",
        "Décris le premier mannequin (2 phrases) : Prénom + Present Continuous + tenue",
        "Décris le deuxième mannequin (2 phrases) : même structure — tenue différente",
        "Conclus en décrivant la réaction du public → The audience is... / The students are...",
        "Compte tes mots (80 à 100) et relis.",
        "Vérifie : am/is/are + verb-ING toujours présents ? Adjectifs avant les noms ?",
    ]
)

# ============================================================
# PARTIE 4 — CORRIGÉS
# ============================================================
add_part_banner(doc, "Section 4 — Corrigés complets")

add_corrige(doc, "CORRIGÉ — Devoir Maison", [
    "TITRE PROPOSÉ : \"A Brilliant Fashion Show at Lycée Moderne de Cocody\"",
    "",
    "This evening, a wonderful fashion show is taking place at Lycée Moderne de Cocody in Abidjan. The students of Terminale are organising the event.",
    "On the runway, Adjoua is walking confidently. She is wearing a long, colourful wax dress with golden embroidery.",
    "Another model, Konan, is showing a fitted blue bazin boubou with silver buttons.",
    "The audience is clapping and cheering loudly. Everyone is enjoying the show. The designers are smiling backstage. It is a great moment for the school!",
    "",
    "TRADUCTION : Ce soir, un magnifique défilé se déroule au Lycée Moderne de Cocody à Abidjan. Les élèves de Terminale organisent l'événement. Sur le podium, Adjoua marche avec assurance. Elle porte une longue robe en wax colorée avec une broderie dorée. Konan présente un boubou en bazin bleu ajusté avec des boutons argentés. Le public applaudit et acclame. Tout le monde apprécie le spectacle. Les stylistes sourient dans les coulisses.",
    "(88 mots — dans les limites requises)",
])
add_corrige(doc, "CORRIGÉ — Exercice 1", [
    "1 → d  (runway = podium / passerelle)",
    "2 → a  (designer = styliste)",
    "3 → e  (audience = public / spectateurs)",
    "4 → c  (outfit = tenue)",
    "5 → b  (backstage = coulisses)",
])
add_corrige(doc, "CORRIGÉ — Exercice 2", [
    "1. dance  → dancing   (verbe en -E muet → retire le E, ajoute -ING)",
    "2. sit    → sitting   (verbe court CVC → double la consonne T, ajoute -ING)",
    "3. show   → showing   (verbe ordinaire → ajoute simplement -ING)",
    "4. write  → writing   (verbe en -E muet → retire le E, ajoute -ING)",
    "5. run    → running   (verbe court CVC → double la consonne N, ajoute -ING)",
])
add_corrige(doc, "CORRIGÉ — Exercice 3", [
    "1. Fatou is wearing a long red dress.  (Fatou porte une longue robe rouge.)",
    "2. The photographer is taking pictures.  (⚠ take → taking, verbe en -E muet)",
    "3. I am looking at the runway.  (Je regarde le podium.)",
    "4. The models are walking on the stage.  (Les mannequins marchent sur la scène.)",
    "5. The designer is showing his new collection.  (Le styliste présente sa nouvelle collection.)",
])
add_corrige(doc, "CORRIGÉ — Exercice 4", [
    "Awa :    \"Awa is wearing a long green wax dress with a golden belt.\"  (Awa porte une longue robe en wax verte avec une ceinture dorée.)",
    "Issiaka : \"Issiaka is wearing a white bazin boubou with blue embroidery.\"  (Issiaka porte un boubou en bazin blanc avec des broderies bleues.)",
    "Nadia :  \"Nadia is wearing a short black skirt with geometric patterns.\"  (Nadia porte une jupe courte noire avec des motifs géométriques.)",
    "✔ Vérification : adjectifs avant le nom ✔ · Present Continuous ✔ · vocabulaire de la tenue ✔",
])
add_corrige(doc, "CORRIGÉ — Exercice 5 (modèle proposé)", [
    "\"Tonight, a wonderful fashion show is taking place at the Palais de la Culture in Treichville, Abidjan.",
    "On the runway, Mariama is walking gracefully. She is wearing a fitted colourful wax dress with golden embroidery.",
    "Another model, Sékou, is showing a long white boubou with blue patterns.",
    "The audience is cheering and clapping loudly. The designer is smiling backstage.",
    "It is a great night for Ivorian fashion!\"",
    "",
    "Vérification : ✔ Present Continuous (is taking place, is walking, is wearing, is showing, is cheering, is clapping, is smiling) · ✔ 2 tenues décrites · ✔ runway, audience, designer, backstage · ✔ Palais de la Culture / Treichville / Abidjan · ✔ Mariama, Sékou",
])

# --- Pied de page final ---
p_final = doc.add_paragraph()
p_final.alignment = WD_ALIGN_PARAGRAPH.CENTER
_para_border(p_final, BLEU, side="top", sz=8)
r(p_final, "Prof. Evelyne ASSI  ·  Anglais LV1  ·  3ème  ·  DPFC/MENET-FP CI · 2025-2026",
  color=GRIS, size=9)
p_quote = doc.add_paragraph()
p_quote.alignment = WD_ALIGN_PARAGRAPH.CENTER
r(p_quote, "\"Speak even if it's not perfect — silence teaches nothing.\"",
  italic=True, color=BLEU, size=10)
r(p_quote, "  (Parle même si c'est imparfait — le silence n'apprend rien.)",
  italic=True, color=GRIS, size=9.5)

# ============================================================
# SAUVEGARDE
# ============================================================
OUTPUT = "LPA_Anglais_3eme_Lecon12.docx"
doc.save(OUTPUT)
print(f"OK — document généré : {OUTPU:q!
T}":q!

