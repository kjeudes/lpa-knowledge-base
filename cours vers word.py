# ============================================================
#  CONVERTISSEUR DE COURS → WORD  (LPA / DPFC Côte d'Ivoire)
#  Google Colab — python-docx  — VERSION 4 (style modèle v2)
#
#  UTILISATION :
#  1. Installe : !pip install python-docx -q
#  2. Colle ton cours dans la variable COURS ci-dessous
#  3. Exécute → téléchargement automatique
#
#  FORMAT DU COURS — BALISES RECONNUES :
#
#  EN-TÊTE (les 2 formats sont acceptés) :
#  ════...════  ligne titre  ════...════     ← en-tête ASCII
#  TITRE_LECON: ...  COMPETENCE: ...        ← méta-lignes optionnelles
#  PROFESSEUR: ...   ORGANISME: ...         ← (auto-extraits de l'en-tête si absents)
#
#  SOMMAIRE (2 formats) :
#  "SOMMAIRE\n  1. Avant la leçon\n  2. La leçon..."   ← ancien format numéroté
#  "SOMMAIRE\n  Section 1 — ...\n    · Phase 1..."      ← nouveau format avec ·
#
#  SECTION 1 — AVANT LA LEÇON
#  ─────────────────────────────────
#  [CAPSULE DE REPRISE — ...]        ← avec ou sans crochets
#   CAPSULE DE REPRISE — ...         ← les deux marchent
#  [ANCRAGE IVOIRIEN]                ← avec ou sans crochets
#   ANCRAGE IVOIRIEN                 ← les deux marchent
#  ▶ HARKNESS N — ...
#
#  SECTION 2 — LA LEÇON
#  Phase N · ...
#  A. TITRE / B. TITRE / C. TITRE    ← sous-sections Phase 2
#  TABLEAU DES STRUCTURES — ...      ← titre de tableau
#  | col | col | col |               ← tableau markdown (auto-converti)
#  ~~~  texte anglais                ← texte de découverte
#  ---TRADUCTION---
#  texte français
#  ~~~
#  ┌─ À RETENIR — ... ─┐ ... └─┘
#  ⚠ ATTENTION — ...
#  FORMATION DU ...                  ← boîte retenir auto
#  ◆ DÉFINITION DPFC — ...
#  ✎ TRACE ÉCRITE
#  ~ ANALOGIE CI 1 / ~ ANALOGIE CI 2
#  5 points essentiels à retenir :   ← déclenche la boîte synthèse
#  SYNTHESE_DEBUT ... SYNTHESE_FIN   ← ou balises explicites
#  ◉ EXERCICE GUIDÉ N — ...
#
#  SECTION 3 — APRÈS LA LEÇON
#  ◎ EXERCICE N — ...
#  ◈ DEVOIR MAISON — ...
#
#  SECTION 4 — CORRIGÉS COMPLETS
#  ✔ CORRIGÉ — ...
#
#  FIN DU COURS :
#  CITATION FINALE                   ← balise optionnelle
#  "Citation en anglais..."
#  (Traduction française)
#  — Auteur
#  Prof. Evelyne ASSI · ... · DPFC/MENET-FP CI · 2025-2026
# ============================================================

COURS = """✅ Français · Leçon 2 — Œuvre intégrale théâtrale · en cours de génération...

---

LPA
Lycée Personnel Autonome

Français · Leçon 2 · Classe de Seconde

L'ŒUVRE INTÉGRALE THÉÂTRALE
L'œuvre intégrale théâtrale

Compétence C1 Lecture · Construction du sens

Professeur Mariam KONATÉ · Coefficient 4
DPFC / MENET-FP Côte d'Ivoire · 2025-2026

════════════════════════════════════════════════════════
  FRANÇAIS · Leçon 2 — L'ŒUVRE INTÉGRALE THÉÂTRALE
  Niveau : Seconde | Série : —
  Professeur : Mariam KONATÉ · Coefficient : 4
  Compétence : C1 Lecture · Construction du sens
  Programme : DPFC/MENET-FP CI 2025-2026
════════════════════════════════════════════════════════

SOMMAIRE
  Section 1 — Avant la leçon
    · Capsule de reprise (Leçon 1 : L'œuvre intégrale narrative)
    · Ancrage ivoirien
    · Questions Harkness 1, 2 et 3
  Section 2 — La leçon
    · Phase 1 — Présentation & prérequis
    · Phase 2 — Découverte guidée
    · Phase 3 — Schémas textuels
    · Phase 4 — Définitions DPFC
    · Phase 5 — Trace écrite + Analogies CI
    · Phase 5b — Synthèse
    · Phase 6 — Exercices guidés (2)
  Section 3 — Après la leçon
    · Exercices 1 à 5 + Devoir maison
  Section 4 — Corrigés complets
    · Corrigé du devoir maison + Exercices 1 à 5

────────────────────────────────────────────────────────
SECTION 1 — AVANT LA LEÇON
────────────────────────────────────────────────────────

  CAPSULE DE REPRISE — Leçon 1 : L'œuvre intégrale narrative
  ────────────────────────────────────────────────────────────

  5 points clés à retenir :
    1 · Le narrateur est la voix qui raconte à l'intérieur du texte — il n'est pas l'auteur
    2 · Le point de vue peut être interne (accès aux pensées), omniscient (sait tout)
        ou externe (observe sans entrer dans les esprits)
    3 · Le schéma narratif suit 5 étapes : SI → EP → P → ER → SF
    4 · La tonalité est l'atmosphère créée par les choix stylistiques de l'auteur
    5 · Lire une œuvre intégrale exige une méthode : observer, relever, interpréter

  3 questions flash :
    Q1 · Quelle est la différence entre l'auteur et le narrateur ?
    R1 · L'auteur écrit le livre (personne réelle) ; le narrateur est la voix fictive
         à l'intérieur du texte qui raconte l'histoire.

    Q2 · Cite les 5 étapes du schéma narratif dans l'ordre.
    R2 · Situation initiale → Élément perturbateur → Péripéties →
         Élément de résolution → Situation finale

    Q3 · Quelle tonalité domine un texte qui exprime des émotions intenses
         à la première personne ?
    R3 · La tonalité lyrique.

  Lien avec la leçon du jour :
    En Leçon 1, tu as appris à lire un RÉCIT : une voix raconte une histoire.
    Aujourd'hui, tu vas découvrir le THÉÂTRE : les personnages parlent EUX-MÊMES,
    directement, sur scène. Il n'y a plus de narrateur — ou presque.
    Le texte de théâtre se lit et se joue différemment du roman.


  ANCRAGE IVOIRIEN
  ─────────────────
  Au Lycée Moderne de Yopougon, le professeur de Français a annoncé que la classe
  va monter une pièce de théâtre pour la fête de fin d'année. Bintou, la déléguée,
  distribue les rôles. Kobenan, timide, reçoit le rôle principal. Il ouvre le texte
  et découvre quelque chose d'étrange : pas de paragraphes, pas de « il dit »,
  pas de « elle répondit ». Juste des noms suivis de deux points, et des répliques.
  « C'est quoi ces didascalies ? » demande-t-il.

  C'est exactement la question à laquelle cette leçon va répondre.
  Lire une pièce de théâtre, ce n'est pas lire un roman :
  c'est lire un texte conçu pour être joué, entendu, vu.


▶ HARKNESS 1 — Qu'est-ce que le théâtre ?
  Q : Quelle est la différence fondamentale entre lire un roman et lire une pièce
      de théâtre ? Pourquoi dit-on que le théâtre est un art « double » ?
  Guide :
    · Le roman est fait pour être lu seul en silence
    · La pièce de théâtre est écrite pour être jouée devant un public
    · Pense à ce que tu vois quand tu lis : dans un roman, tu imagines tout ;
      au théâtre, tout doit être montré sur scène
  Corrigé :
    Le roman est un art de la lecture solitaire : le narrateur guide le lecteur,
    décrit les lieux, entre dans les pensées des personnages. La pièce de théâtre,
    elle, est un art double : elle existe d'abord comme texte écrit (le texte
    dramatique), puis comme spectacle vivant (la représentation scénique). L'auteur
    dramatique ne peut pas décrire les pensées de ses personnages directement —
    tout doit passer par la parole et le geste. C'est pourquoi le théâtre est
    souvent considéré comme l'art le plus exigeant pour l'écrivain : convaincre
    sans expliquer, montrer sans raconter.

▶ HARKNESS 2 — Le conflit dramatique
  Q : Dans toute pièce de théâtre, il y a un conflit. Pourquoi un conflit est-il
      indispensable au théâtre ? Qu'est-ce qui rend un conflit dramatique
      intéressant pour le spectateur ?
  Guide :
    · Sans conflit, il n'y a pas d'action, donc pas de pièce
    · Un conflit oppose deux forces, deux volontés, deux visions du monde
    · Pense à un débat entre toi et quelqu'un de ta famille : c'est déjà du théâtre !
  Corrigé :
    Le conflit est le moteur du théâtre. Sans opposition entre personnages —
    ou entre un personnage et lui-même — il n'y a pas de tension, donc pas
    d'intérêt dramatique. Le conflit peut être extérieur (deux personnages
    s'affrontent pour le pouvoir, l'amour, la survie) ou intérieur (un personnage
    déchiré entre deux choix). Ce qui rend un conflit intéressant, c'est que le
    spectateur ne sait pas comment il va se résoudre : il attend, il s'inquiète,
    il s'identifie. C'est ce qu'on appelle la tension dramatique, et c'est elle
    qui maintient l'attention du public du début à la fin.

▶ HARKNESS 3 — Théâtre africain et théâtre classique
  Q : Le théâtre existe en Afrique depuis longtemps sous des formes variées
      (masques, cérémonies, contes joués). En quoi le théâtre écrit africain
      moderne hérite-t-il de ces traditions tout en s'en distinguant ?
  Guide :
    · Pense aux cérémonies de la communauté, aux griots qui jouent des personnages
    · Le théâtre écrit africain utilise les codes européens (actes, scènes, didascalies)
    · Mais il intègre des langues, des valeurs, des conflits propres à l'Afrique
  Corrigé :
    En Afrique, le théâtre existait bien avant la colonisation : cérémonies de masque,
    contes joués par les griots, rites d'initiation dramatisés. Ces formes mettaient
    en scène la communauté, ses conflits et ses valeurs. Le théâtre africain écrit
    moderne — comme celui de Bernard Dadié, d'Aimé Césaire ou de Sony Labou Tansi —
    hérite de cette tradition de la parole vivante et collective, tout en adoptant
    la structure du théâtre occidental (actes, scènes, répliques, didascalies).
    Il y a donc une hybridation : les formes sont européennes, mais les voix,
    les enjeux et les personnages sont profondément africains.


────────────────────────────────────────────────────────
SECTION 2 — LA LEÇON
────────────────────────────────────────────────────────

Phase 1 · Présentation / Prérequis
───────────────────────────────────

Titre : L'œuvre intégrale théâtrale
Objectifs :
  · Identifier les composantes spécifiques d'un texte de théâtre
    (didascalies, répliques, actes, scènes, conflit dramatique)
  · Reconnaître les tonalités théâtrales dominantes
    (tragique, comique, dramatique)
  · Lire une pièce en appliquant une méthode d'analyse de scène

Prérequis nécessaires :
  · Connaître la notion de narrateur et de point de vue (Leçon 1)
  · Savoir identifier une situation initiale et un conflit dans un récit
  · Avoir une idée générale de ce qu'est une pièce de théâtre

┌─ À RETENIR ──────────────────────────────────────────┐
  Au théâtre, il n'y a pas de narrateur qui raconte :
  les personnages parlent directement, en scène.
  L'auteur doit tout montrer par la parole et le geste.
  "En vue du BAC, retenez que l'analyse de scène
  théâtrale est distincte de l'analyse narrative :
  on ne cherche pas qui raconte, mais comment les
  personnages s'affrontent et ce que cela révèle."
└──────────────────────────────────────────────────────┘

Œuvre de référence pour cette leçon :
  Monsieur Thôgô-gnini · Bernard Binlin Dadié · 1970 · Éditions Présence Africaine
  Auteur : Ivoirien · Langue : Français · Contexte : Afrique postcoloniale
  Thèmes : colonialisme · trahison · pouvoir · identité · satire sociale
  Note : pièce majeure de Bernard Dadié, auteur ivoirien emblématique,
         ancien directeur des Affaires culturelles de Côte d'Ivoire.
         Régulièrement étudiée dans les lycées ivoiriens.


Phase 2 · Découverte guidée
────────────────────────────

  I. LA STRUCTURE D'UN TEXTE DE THÉÂTRE

  Un texte de théâtre — appelé aussi texte dramatique — est organisé d'une façon
  très particulière, différente du roman.

  Il est découpé en ACTES et en SCÈNES :

  L'ACTE
  → Grande division de la pièce, comparable à un chapitre de roman
  → Indique une étape importante de l'action
  → Une pièce peut avoir 1, 3 ou 5 actes selon les traditions

  LA SCÈNE
  → Subdivision à l'intérieur d'un acte
  → Change chaque fois qu'un personnage entre ou sort de scène
  → Numérotée : Scène 1, Scène 2, Scène 3…

  Exemple de présentation dans un texte :
  ACTE II · Scène 3
  (Thôgô-gnini entre en courant. Il croise Bénié près de la porte.)
  THÔGÔ-GNINI · Tout est perdu, tout !
  BÉNIÉ · Calme-toi. Personne ne nous a vus.

  ┌─ À RETENIR ──────────────────────────────────────────┐
    Structure d'une pièce = Actes → Scènes → Répliques
    · Acte = grande étape de l'action
    · Scène = moment délimité par les entrées/sorties
    · Réplique = parole prononcée par un personnage
  └──────────────────────────────────────────────────────┘


  II. LES DIDASCALIES

  Les didascalies sont les indications que l'auteur donne en dehors des paroles
  des personnages. Elles sont généralement écrites en italique ou entre parenthèses.

  Elles peuvent indiquer :
  · Les déplacements des personnages sur scène
    → Exemple : (Il s'avance vers la fenêtre.)
  · Les gestes et expressions
    → Exemple : (Il rit, puis se fige.)
  · Le décor, les costumes, l'éclairage
    → Exemple : (La scène représente un bureau colonial surchargé d'objets.)
  · Le ton à adopter pour une réplique
    → Exemple : (Avec mépris.) ou (À voix basse.)

  ⚠ ATTENTION
    Les didascalies ne sont PAS prononcées sur scène.
    Le comédien les lit pour savoir comment jouer,
    mais le spectateur ne les entend jamais.
    En analyse de texte, les didascalies sont pourtant
    essentielles : elles révèlent l'intention de l'auteur.


  III. LA RÉPLIQUE ET LE DIALOGUE DRAMATIQUE

  La réplique est l'unité de base du théâtre : c'est ce qu'un personnage dit.
  L'ensemble des répliques entre plusieurs personnages forme le DIALOGUE DRAMATIQUE.

  Caractéristiques du dialogue dramatique au théâtre :
  · Il est direct — pas de guillemets, pas de « dit-il »
  · Il est le seul moyen pour les personnages de se révéler
  · Il porte en lui le conflit (les personnages ne sont pas toujours d'accord)
  · Il peut contenir des sous-entendus (ce qu'on dit ≠ ce qu'on pense)

  Formes particulières de parole au théâtre :

  Le MONOLOGUE
  → Un personnage seul sur scène qui parle à voix haute
  → Révèle ses pensées intérieures — c'est le seul moment où le théâtre ressemble
    à un roman avec point de vue interne
  → Exemple : Hamlet qui dit « Être ou ne pas être… »

  La TIRADE
  → Longue réplique d'un personnage qui parle longuement sans être interrompu
  → Souvent utilisée pour argumenter, persuader ou révéler une vérité

  L'APARTÉ
  → Le personnage dit quelque chose que le public entend mais que les autres
    personnages sur scène sont censés ne pas entendre
  → Crée un effet de complicité avec le spectateur

  ┌─ À RETENIR ──────────────────────────────────────────┐
    Les 3 formes spéciales de parole au théâtre :
    · Monologue = seul sur scène, pense à voix haute
    · Tirade = longue réplique argumentée
    · Aparté = dit au public, pas aux autres personnages
  └──────────────────────────────────────────────────────┘


  IV. LE CONFLIT DRAMATIQUE

  Le CONFLIT DRAMATIQUE est l'opposition centrale qui crée la tension de la pièce.
  Sans conflit, il n'y a pas de théâtre.

  Types de conflits :

  Conflit EXTÉRIEUR (entre personnages)
  → Deux personnages s'affrontent pour des intérêts opposés
  → Exemple dans Thôgô-gnini : le personnage trahit sa communauté pour s'enrichir
    avec les colons — conflit entre lui et ceux qu'il exploite

  Conflit INTÉRIEUR (dans un personnage)
  → Un personnage est déchiré entre deux choix, deux valeurs
  → Exemple : choisir entre la fidélité à son peuple et l'appât du gain

  Conflit SOCIAL ou POLITIQUE
  → La pièce met en scène une opposition entre des classes, des groupes, des systèmes
  → Thôgô-gnini est une satire du collaborationnisme colonial : conflit entre dominants
    et dominés, entre tradition et corruption

  ⚠ ATTENTION
    Le conflit n'est pas toujours violent ou spectaculaire.
    Il peut se manifester dans un simple dialogue : deux personnages
    qui ne sont pas d'accord créent déjà du conflit dramatique.


  V. LES TONALITÉS THÉÂTRALES

  Les principales tonalités au théâtre :

  Tonalité TRAGIQUE
  → Le personnage est écrasé par une fatalité qu'il ne peut éviter
  → Vocabulaire de la mort, de la chute, de l'inévitable
  → Exemple : un personnage qui sait qu'il va mourir et accepte son destin

  Tonalité COMIQUE
  → Vise à provoquer le rire par des situations absurdes, des malentendus,
    des personnages ridicules
  → Types : comique de mots · comique de situation · comique de caractère
  → Exemple dans Thôgô-gnini : le personnage principal est ridiculisé par
    sa propre cupidité — comique de caractère

  Tonalité DRAMATIQUE
  → Action tendue, événements violents ou graves, rythme rapide
  → Pas forcément tragique : le dénouement peut être heureux
  → Exemple : une scène de confrontation violente où le sort d'un personnage
    est en jeu

  Tonalité POLÉMIQUE / SATIRIQUE
  → Critique sociale ou politique, souvent par l'ironie
  → Thôgô-gnini : Dadié dénonce le collaborationnisme par la satire —
    le personnage est grotesque pour mieux être condamné


Phase 3 · Schémas textuels
────────────────────────────

[SCHÉMA 1 — Structure d'un texte de théâtre]
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
Description visuelle :
  · Forme générale : arbre hiérarchique vertical de haut en bas
  · Niveau 1 (haut, rectangle rouge) : "LA PIÈCE DE THÉÂTRE"
  · Niveau 2 (deux rectangles bleus, côte à côte) : "ACTE I" | "ACTE II" | "ACTE III…"
  · Niveau 3 (sous chaque acte, rectangles verts) : "Scène 1" | "Scène 2" | "Scène 3"
  · Niveau 4 (sous chaque scène, rectangles jaunes) : "Réplique" | "Didascalie" | "Monologue"
  · Flèches descendantes reliant chaque niveau au suivant
  · Légende à droite :
    Rouge = structure globale
    Bleu = grande division
    Vert = subdivision
    Jaune = unité de base
Note générateur : draw.io (organigramme vertical) ou Canva (modèle "Hiérarchie")
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

[SCHÉMA 2 — Carte mentale : Les outils d'analyse théâtrale]
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
Description visuelle :
  · Nœud central (ovale, fond bordeaux) : "TEXTE DE THÉÂTRE"
  · 4 branches principales rayonnant du centre :
    Branche 1 (fond bleu, haut gauche) : "STRUCTURE"
      → Sous-branches : "Actes" · "Scènes" · "Répliques"
    Branche 2 (fond vert, haut droite) : "PAROLE"
      → Sous-branches : "Dialogue" · "Monologue" · "Tirade" · "Aparté"
    Branche 3 (fond orange, bas gauche) : "DIDASCALIES"
      → Sous-branches : "Gestes" · "Décor" · "Ton" · "Déplacements"
    Branche 4 (fond violet, bas droite) : "CONFLIT"
      → Sous-branches : "Extérieur" · "Intérieur" · "Social/Politique"
  · Style : lignes courbées, couleurs vives
Note générateur : XMind ou Canva (modèle "Mind Map")
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━


Phase 4 · Définitions DPFC
────────────────────────────

◆ DÉFINITION DPFC — Texte dramatique
  Texte écrit pour être représenté sur scène. Il se compose de répliques
  (paroles des personnages) et de didascalies (indications scéniques de l'auteur).
  Il s'oppose au texte narratif en ce que les personnages parlent directement,
  sans l'intermédiaire d'un narrateur.

◆ DÉFINITION DPFC — Didascalie
  Indication donnée par l'auteur dramatique en dehors des répliques, généralement
  en italique ou entre parenthèses. Elle précise les gestes, déplacements, décors,
  costumes, éclairages ou le ton d'une réplique. Elle est lue par les acteurs
  et le metteur en scène, mais jamais prononcée sur scène.

◆ DÉFINITION DPFC — Acte
  Grande division d'une pièce de théâtre, comparable à un chapitre dans un roman.
  Un acte correspond à une étape importante de l'action dramatique. Une pièce
  peut comporter un ou plusieurs actes.

◆ DÉFINITION DPFC — Scène
  Subdivision d'un acte, délimitée par les entrées et sorties des personnages.
  Chaque fois qu'un personnage entre ou quitte la scène, une nouvelle scène commence.

◆ DÉFINITION DPFC — Conflit dramatique
  Opposition centrale entre des forces, des volontés ou des valeurs contraires,
  qui crée la tension de la pièce et maintient l'attention du spectateur.
  Le conflit peut être extérieur (entre personnages), intérieur (dans un personnage)
  ou social/politique (entre groupes ou systèmes).

◆ DÉFINITION DPFC — Monologue
  Discours d'un personnage seul sur scène, qui exprime ses pensées à voix haute.
  Forme la plus proche du point de vue interne du roman : elle révèle la vie
  intérieure du personnage.

◆ DÉFINITION DPFC — Tirade
  Longue réplique d'un personnage, sans interruption des autres.
  Elle sert généralement à argumenter, persuader, ou révéler une vérité importante.

◆ DÉFINITION DPFC — Aparté
  Réplique qu'un personnage prononce à l'intention du public, sans que les autres
  personnages soient censés l'entendre. Crée une complicité entre le personnage
  et le spectateur.

◆ DÉFINITION DPFC — Tonalité tragique
  Atmosphère d'une pièce dans laquelle le personnage est écrasé par une fatalité
  inévitable. Elle se manifeste par le vocabulaire de la mort, de la chute,
  et de l'inéluctable.

◆ DÉFINITION DPFC — Tonalité comique
  Atmosphère visant à provoquer le rire du spectateur. Elle peut reposer sur
  le comique de mots (jeux de langage), de situation (malentendus, quiproquos)
  ou de caractère (personnages ridicules ou excessifs).


Phase 5 · Trace écrite + Analogies CI
───────────────────────────────────────

✎ TRACE ÉCRITE (à recopier)

  LEÇON 2 · L'ŒUVRE INTÉGRALE THÉÂTRALE
  Français · Seconde · Prof. Mariam KONATÉ

  I. STRUCTURE D'UN TEXTE DE THÉÂTRE
  · Pièce → Actes → Scènes → Répliques / Didascalies
  · Didascalie = indication scénique de l'auteur (non prononcée sur scène)

  II. LES FORMES DE PAROLE AU THÉÂTRE
  · Dialogue dramatique : échange entre personnages — porte le conflit
  · Monologue : seul sur scène — révèle les pensées intérieures
  · Tirade : longue réplique argumentée
  · Aparté : dit au public, pas aux autres personnages

  III. LE CONFLIT DRAMATIQUE
  · Extérieur : entre personnages (intérêts opposés)
  · Intérieur : dans un personnage (déchiré entre deux choix)
  · Social / Politique : entre groupes, systèmes, valeurs

  IV. LES TONALITÉS AU THÉÂTRE
  · Tragique : fatalité, mort inévitable
  · Comique : rire — de mots / de situation / de caractère
  · Dramatique : tension, action grave
  · Satirique : critique par l'ironie et le ridicule

  ŒUVRE DE RÉFÉRENCE : Monsieur Thôgô-gnini · Bernard Binlin Dadié · 1970
  → Tonalités dominantes : comique + satirique (critique du collaborationnisme)
  → Conflit : extérieur (Thôgô-gnini vs sa communauté) + social (dominants/dominés)


~ ANALOGIE CI 1
  Le dialogue dramatique ressemble à une palabre dans une cour de quartier
  à Treichville ou Adjamé. Chacun prend la parole à son tour. Les tensions
  montent. On ne dit pas toujours ce que l'on pense vraiment — les sous-entendus
  fusent. Et le « chef de palabre » (le metteur en scène) décide comment
  tout cela va se dérouler.
  La palabre africaine est déjà une forme de théâtre.

~ ANALOGIE CI 2
  Les didascalies, c'est comme les instructions du coach de l'ASEC Mimosas
  écrites sur son tableau avant un match : les joueurs (acteurs) les lisent,
  les intègrent, mais le public, lui, ne voit que le jeu.
  Les didascalies guident les acteurs ; le spectateur n'en voit que le résultat.


Phase 5b · Synthèse
─────────────────────

5 POINTS ESSENTIELS À RETENIR :
  1 · Au théâtre, il n'y a pas de narrateur : les personnages parlent directement
  2 · La didascalie est une instruction de l'auteur — elle n'est jamais prononcée
  3 · La pièce est structurée en actes, puis en scènes, puis en répliques
  4 · Le conflit dramatique est le moteur de toute pièce de théâtre
  5 · Les tonalités dominantes d'une pièce se justifient par des relevés précis

5 MOTS CLÉS + DÉFINITION COURTE :
  · Didascalie      : indication scénique de l'auteur, hors répliques
  · Acte / Scène    : grandes et petites divisions de la pièce
  · Monologue       : parole d'un personnage seul, révèle ses pensées
  · Conflit dramatique : opposition qui crée la tension de la pièce
  · Tonalité comique : atmosphère qui vise à provoquer le rire

3 ERREURS FRÉQUENTES + CORRECTION :
  Erreur 1 · Lire les didascalies comme si elles étaient des répliques
    → Correction : les didascalies ne sont jamais prononcées sur scène.
      Ce sont des instructions pour les comédiens et le metteur en scène.

  Erreur 2 · Confondre monologue et tirade
    → Correction : le monologue = personnage SEUL sur scène.
      La tirade = longue réplique, mais d'autres personnages peuvent être présents.

  Erreur 3 · Dire « la pièce est comique » sans justifier
    → Correction : il faut préciser le TYPE de comique (mots / situation / caractère)
      ET citer un exemple précis du texte.


Phase 6 · Exercices guidés
────────────────────────────

◉ EXERCICE GUIDÉ 1 — Analyser une scène de théâtre · Notion ciblée : didascalies, répliques, conflit

  Énoncé :
  Lis cet extrait de théâtre (inspiré de Monsieur Thôgô-gnini · Bernard Dadié) :

  ACTE I · Scène 2
  (La case de Thôgô-gnini. Décor sobre : un tabouret, une natte roulée, une lampe à huile.
  Thôgô-gnini est assis, comptant des pièces de monnaie. Bénié entre précipitamment.)

  BÉNIÉ · (hors d'haleine) Ils arrivent ! Les anciens du village sont là.
  THÔGÔ-GNINI · (sans lever les yeux) Qu'ils attendent.
  BÉNIÉ · Mais… tu les as convoqués toi-même !
  THÔGÔ-GNINI · (avec un sourire ironique) Et alors ? Celui qui attend apprend
                 la patience. Moi, j'apprends à compter.
  BÉNIÉ · (à voix basse, pour lui-même) Il a vendu son âme pour quelques pièces.

  Questions :
  A · Identifiez et listez toutes les didascalies de cet extrait.
  B · Quelle est la réplique de Bénié qui ressemble à un aparté ? Justifiez.
  C · Quel conflit est en train de naître dans cet extrait ?
  D · Quelle tonalité domine ? Justifiez par un relevé.

  MÉTHODE
  Étape 1 · Repère tout ce qui est entre parenthèses ou en italique
    → (La case de Thôgô-gnini. Décor sobre…) → didascalie de décor
    → (hors d'haleine) → didascalie de jeu (état physique de Bénié)
    → (sans lever les yeux) → didascalie de geste (attitude de Thôgô-gnini)
    → (avec un sourire ironique) → didascalie de ton
    → (à voix basse, pour lui-même) → didascalie + indication d'aparté
  Étape 2 · Identifie l'aparté
    → Bénié dit : « Il a vendu son âme pour quelques pièces »
    → (à voix basse, pour lui-même) → il ne s'adresse pas à Thôgô-gnini
    → C'est un aparté : il parle au public, pas au personnage
  Étape 3 · Identifie le conflit
    → Thôgô-gnini méprise les anciens (« Qu'ils attendent »)
    → Conflit extérieur naissant entre Thôgô-gnini et sa communauté
    → Conflit de valeurs : cupidité (les pièces) vs respect traditionnel (les anciens)
  Étape 4 · Identifie la tonalité
    → « avec un sourire ironique » + « apprendre à compter » → ironie mordante
    → Le personnage est ridiculisé par sa propre arrogance
    → Tonalité satirique (et légèrement comique de caractère)

  ✔ Vérification : chaque réponse est justifiée par un relevé dans le texte. Oui.
  ✔ Conclusion :
    A · 5 didascalies : décor, état de Bénié, geste de Thôgô-gnini, ton ironique, aparté de Bénié
    B · Aparté : « Il a vendu son âme pour quelques pièces » — dit pour lui-même, donc pour le public
    C · Conflit entre Thôgô-gnini et les anciens du village — conflit de valeurs (cupidité vs respect)
    D · Tonalité satirique — justifiée par « sourire ironique » et le ridicule du personnage
  ✔ Ce que cet exercice mobilise :
    · Identification des didascalies dans un texte de théâtre
    · Distinction réplique / aparté
    · Repérage du conflit dramatique
    · Identification et justification d'une tonalité


◉ EXERCICE GUIDÉ 2 — Le monologue · Notion ciblée : monologue, pensée intérieure, tonalité

  Énoncé :
  Lis ce monologue fictif inspiré du théâtre africain contemporain :

  (Seul sur scène. Nuit. Une bougie vacille. Kouamé s'avance lentement vers le public.)

  KOUAMÉ · Vingt ans que je travaille pour eux. Vingt ans. Et ce soir, ils me
            disent que je suis trop vieux, trop lent, trop… africain. (Il rit
            amèrement.) Trop africain ! Dans mon propre pays ! Mes enfants
            attendent à la maison. Ma femme a les yeux rouges depuis ce matin.
            (Il se retourne, comme s'il voyait quelqu'un derrière lui.)
            Mon père m'avait dit : "Ne vends jamais ta dignité pour un salaire."
            J'aurais dû l'écouter. (Long silence.) Que faire maintenant ?

  Questions :
  A · Pourquoi s'agit-il d'un monologue ? Citez un indice dans le texte.
  B · Quelle est la fonction de ce monologue dans la pièce ?
  C · Identifiez deux tonalités présentes. Justifiez chacune par un relevé.
  D · En quoi ce monologue ressemble-t-il au point de vue interne du roman ?

  MÉTHODE
  Étape 1 · Vérifie les conditions du monologue
    → Didascalie d'ouverture : « Seul sur scène » → confirme le monologue
    → Kouamé parle sans interlocuteur visible — il se parle à lui-même
  Étape 2 · Identifie la fonction du monologue
    → Il révèle les pensées intérieures de Kouamé (injustice ressentie, souvenir du père)
    → Il informe le spectateur sur la situation (licenciement, famille, passé)
    → Il crée de l'empathie : le public comprend sa douleur
  Étape 3 · Identifie les tonalités
    → Tonalité PATHÉTIQUE : « Ma femme a les yeux rouges » · « Que faire maintenant ? »
      — appel à la compassion du spectateur
    → Tonalité POLÉMIQUE / SATIRIQUE : « trop africain ! Dans mon propre pays ! »
      — critique implicite de la discrimination postcoloniale
  Étape 4 · Compare avec le point de vue interne du roman
    → Dans un roman en point de vue interne, le narrateur entre dans la tête du personnage
    → Dans ce monologue, Kouamé dit tout ce qu'il pense à voix haute — même effet
    → Différence : au théâtre, c'est dit devant un public, pas narré en silence

  ✔ Vérification : fonction identifiée, deux tonalités justifiées, comparaison argumentée. Oui.
  ✔ Conclusion :
    A · Monologue confirmé par « Seul sur scène » + aucun interlocuteur dans les répliques
    B · Révèle les pensées intérieures · informe le spectateur · crée de l'empathie
    C · Pathétique (« Ma femme a les yeux rouges ») + Polémique (« trop africain ! »)
    D · Comme le point de vue interne, le monologue donne accès direct aux pensées du personnage
  ✔ Ce que cet exercice mobilise :
    · Identification et justification du monologue
    · Analyse de la fonction dramatique d'un monologue
    · Identification et justification de deux tonalités
    · Lien entre théâtre et roman (point de vue interne / monologue)


────────────────────────────────────────────────────────
SECTION 3 — APRÈS LA LEÇON
────────────────────────────────────────────────────────

◎ EXERCICE 1 — Identifier les éléments d'un texte de théâtre
  Notions travaillées : didascalies, répliques, actes et scènes

  Lis cet extrait :

  ACTE II · Scène 1
  (Une rue d'Abidjan. Il fait chaud. Soro et Diallo se croisent devant une boutique.)

  SORO · (surpris) Diallo ! Ça fait combien d'années ?
  DIALLO · (froidement) Trop longtemps. Ou pas assez.
  SORO · (déconcerté) Qu'est-ce que tu veux dire par là ?
  DIALLO · (se retournant pour partir) Tu le sais très bien.
  (Il sort. Soro reste seul, immobile.)

  Questions :
  A · Combien de didascalies trouve-t-on dans cet extrait ? Listez-les.
  B · Combien de répliques y a-t-il ?
  C · À quel acte et à quelle scène appartient cet extrait ?
  D · Quel sentiment domine dans cet extrait ? Justifiez par un relevé.

  GUIDE
  Étape 1 · Repère tout ce qui est entre parenthèses (= didascalies) et compte-les
  Étape 2 · Compte les prises de parole des personnages (une par personnage par tour)
  Étape 3 · Lis l'en-tête de la scène
  Étape 4 · Lis le ton des répliques et les didascalies de ton pour identifier le sentiment


◎ EXERCICE 2 — Distinguer les formes de parole au théâtre
  Notions travaillées : dialogue, monologue, tirade, aparté

  Pour chacune des situations suivantes, indique s'il s'agit d'un dialogue,
  d'un monologue, d'une tirade ou d'un aparté. Justifie ta réponse.

  Situation A :
  (Seul sur scène.) ATTA · Qu'ai-je fait ? Pourquoi ai-je trahi mon propre frère ?
  Je ne mérite pas ce que j'ai reçu. (Il s'effondre sur le sol.)

  Situation B :
  YABA · (à voix basse, tournée vers le public, pendant que les autres parlent entre eux)
  Il ne sait pas que je connais son secret.

  Situation C :
  KODJO · Je veux te faire comprendre une chose importante. Depuis le début de cette affaire,
  j'ai tout supporté en silence : les mensonges, les trahisons, les humiliations.
  Mais cette nuit, c'est terminé. Je reprends ma parole et ma dignité.
  À partir de demain, les choses changeront, que tu le veuilles ou non.

  Situation D :
  ASSA · Tu as vu ce qui s'est passé ?
  MOUSSA · Non, raconte-moi.
  ASSA · Il a tout refusé. Tout.
  MOUSSA · Même l'argent ?
  ASSA · Même l'argent.

  GUIDE
  Étape 1 · Vérifie si le personnage est seul ou non sur scène
  Étape 2 · Vérifie si la réplique est courte ou longue
  Étape 3 · Vérifie si d'autres personnages sont censés entendre ce qui est dit
  Étape 4 · Nomme la forme et justifie avec un détail du texte


◎ EXERCICE 3 — Identifier le conflit dramatique
  Notions travaillées : conflit extérieur, intérieur, social

  Lis ces trois situations de théâtre. Pour chacune, identifie le TYPE de conflit
  (extérieur / intérieur / social ou politique) et explique-le en deux phrases.

  Situation 1 :
  Dans une pièce, un père veut marier sa fille à un homme riche.
  La fille aime un jeune homme sans argent. Ils s'affrontent à chaque scène.

  Situation 2 :
  Un fonctionnaire découvre que son chef est corrompu.
  Il doit choisir : dénoncer son chef et perdre son emploi,
  ou se taire et garder ses avantages.

  Situation 3 :
  Dans une pièce africaine postcoloniale, des ouvriers s'organisent pour
  réclamer leurs droits face à un patron qui les exploite.

  GUIDE
  Étape 1 · Identifie qui s'oppose à qui (ou quoi s'oppose à quoi)
  Étape 2 · Détermine si le conflit est entre personnes, dans une personne,
             ou entre groupes / systèmes
  Étape 3 · Nomme le type et explique en deux phrases


◎ EXERCICE 4 — Les tonalités au théâtre
  Notions travaillées : tragique, comique, dramatique, satirique

  Lis ces deux extraits et réponds aux questions.

  Extrait A :
  TOGBA · (en larmes) Mon fils… ils l'ont emmené cette nuit. Sans un mot.
           Sans un adieu. Je n'entendrai plus jamais sa voix.
  LA VOISINE · Courage, Togba. Le ciel voit tout.
  TOGBA · Le ciel voit tout… mais il ne fait rien. (Il baisse la tête.)

  Extrait B :
  MONSIEUR KASSI · (bombant le torse devant un miroir) Avec ce costume,
                    je ressemble à un ministre ! Non — à un président !
  SA FEMME · (en coulisses) Kassi ! Tu as encore oublié d'enlever l'étiquette du prix !
  MONSIEUR KASSI · (se figeant) Quelle étiquette…?
                   (Il se retourne et tire sur l'étiquette qui pendait dans son dos.)
                   Par le diable !

  Questions :
  A · Quelle est la tonalité dominante de l'Extrait A ? Justifiez par deux relevés.
  B · Quelle est la tonalité dominante de l'Extrait B ? Quel type de comique est en jeu ?
  C · En quoi ces deux extraits montrent-ils que le théâtre peut traiter de sujets
      très différents avec des effets opposés sur le spectateur ?

  GUIDE
  Étape 1 · Note le sentiment que chaque extrait provoque en toi (tristesse / rire)
  Étape 2 · Relève les mots ou didascalies qui créent ce sentiment
  Étape 3 · Nomme la tonalité et justifie
  Étape 4 · Pour la question C, réfléchis à la diversité des fonctions du théâtre


◎ EXERCICE 5 — Analyse complète d'une scène
  Notions travaillées : ensemble des outils de la leçon

  Lis cette scène intégrale :

  ACTE III · Scène 4
  (La case du chef. Soir. Une lampe éclaire faiblement la pièce. Le chef GNANDÉ est assis,
  entouré de deux conseillers. Entre BOUA, chapeau à la main, légèrement courbé.)

  BOUA · Chef, je viens pour l'affaire du terrain.
  GNANDÉ · (sans le regarder) Quel terrain ?
  BOUA · Vous le savez bien, chef. Celui que mon père a cultivé pendant trente ans.
  GNANDÉ · (aux conseillers, avec un sourire) Il parle de son père. (Aux conseillers.)
             Et son père, où est-il maintenant ?
  CONSEILLER 1 · (doucement) Il est décédé, chef.
  GNANDÉ · (avec fausse compassion) Ah… Mes condoléances. (Pause.) Mais les morts
             n'ont pas de terrain. Les terrains appartiennent aux vivants. (Il rit.)
  BOUA · (serrant son chapeau, contenant sa colère) Chef… ce terrain nourrit
          mes six enfants.
  GNANDÉ · (se levant, mettant fin à l'entretien) Six enfants. Voilà un problème
             de famille. Pas un problème de terrain. Bonne soirée, Boua.
  (Boua sort lentement. La porte se referme. Seul un conseiller reste.)
  BOUA · (dans l'embrasure, à voix basse, pour lui-même) Ce n'est pas fini.

  Questions :
  A · Identifiez au moins 4 didascalies et précisez leur type (décor / geste / ton…).
  B · Quelle forme de parole Boua utilise-t-il à la dernière réplique ? Justifiez.
  C · Décrivez le conflit dramatique de cette scène : qui s'oppose à qui ? Pour quoi ?
  D · Identifiez deux tonalités et justifiez chacune par un relevé.
  E · Comment la didascalie « serrant son chapeau, contenant sa colère » enrichit-elle
      notre compréhension du personnage de Boua ?

  GUIDE
  Étape 1 · Liste tout ce qui est entre parenthèses et classe chaque didascalie par type
  Étape 2 · Repère la dernière réplique de Boua et vérifie la présence d'autres personnages
  Étape 3 · Identifie les deux camps en opposition et l'enjeu de leur affrontement
  Étape 4 · Lis le ton général (grave ? ironique ? drôle ?) et relève deux indices textuels
  Étape 5 · Analyse ce que le geste révèle sur l'état intérieur de Boua
             (ce qu'il ne peut pas dire mais montre physiquement)


◈ DEVOIR MAISON — Écrire une scène de théâtre
  Durée conseillée : 50 minutes

  Énoncé :
  Écris une scène de théâtre courte (15 à 20 répliques) dans laquelle deux personnages
  s'affrontent autour d'un conflit que tu choisiras librement.
  Le contexte doit être ivoirien ou africain.

  Contrainte 1 · Présenter la scène correctement : ACTE [N] · Scène [N] + didascalie
                 d'ouverture décrivant le lieu et les personnages présents
  Contrainte 2 · Écrire au minimum 4 didascalies réparties dans la scène
                 (ton, geste, déplacement, état physique)
  Contrainte 3 · Inclure au moins une tirade (5 répliques consécutives d'un même personnage)
  Contrainte 4 · Terminer par un aparté du personnage qui « perd » la confrontation
  Contrainte 5 · Créer une tonalité dominante claire (dramatique, tragique ou satirique)
                 et la maintenir par le choix des mots et des didascalies

  GUIDE (sans corrigé)
  Étape 1 · Choisis ton conflit : une dispute familiale, une injustice au travail,
             un affrontement politique, une trahison entre amis
  Étape 2 · Nomme tes deux personnages et décide qui « gagne » et qui « perd »
  Étape 3 · Écris la didascalie d'ouverture : lieu, heure, ambiance, personnages
  Étape 4 · Rédige le dialogue en faisant monter la tension progressivement
  Étape 5 · Place la tirade au moment le plus intense du conflit
  Étape 6 · Rédige l'aparté final avec la didascalie appropriée
  Étape 7 · Relis et vérifie les 5 contraintes


────────────────────────────────────────────────────────
SECTION 4 — CORRIGÉS COMPLETS
────────────────────────────────────────────────────────

✔ CORRIGÉ — DEVOIR MAISON
  (Exemple de production conforme — à titre indicatif)

  ACTE I · Scène 1
  (La boutique d'un commerçant au marché de Bouaké. Fin d'après-midi.
  Des sacs de vivriers s'empilent. KANGA, le patron, est assis derrière un comptoir.
  Entre YÉO, son employé depuis dix ans, tenant une lettre à la main.)

  YÉO · (posant la lettre sur le comptoir) Patron, j'ai reçu ça ce matin.
  KANGA · (sans regarder) Quoi encore ?
  YÉO · C'est une lettre de licenciement. Avec votre signature.
  KANGA · (levant les yeux, froidement) Et alors ?
  YÉO · Dix ans, Kanga. Dix ans que je suis là à cinq heures du matin.
  KANGA · Les affaires, c'est les affaires. J'ai trouvé quelqu'un de moins cher.
  YÉO · (élevant la voix) Quelqu'un de moins cher ! Mes enfants ont mangé grâce
          à cette boutique. Ma mère est allée à l'hôpital grâce à ce travail.
          Vous ne pouvez pas effacer dix ans d'un trait de plume.
          Vous m'avez vu grandir ici. Vous connaissez ma famille.
          Est-ce que ça ne compte pour rien ?
  KANGA · (impassible) Ça compte pour de bons souvenirs. Pas pour de l'argent.
  YÉO · (stupéfait) Vous n'avez pas honte ?
  KANGA · (se levant) Honte ? J'ai une famille à nourrir, moi aussi.
           Ta remplaçante commence lundi. Je t'ai fait un bon de sortie.
           (Il pousse un papier vers lui.)
  YÉO · (sans le prendre, les mains tremblantes) Je ne signerai pas ça.
  KANGA · (haussant les épaules) Tu n'as pas le choix.
  (Un long silence. Yéo prend le papier, le regarde, puis lève les yeux vers Kanga.)
  YÉO · (à voix basse, se retournant vers le public) Il croit avoir gagné.
          Mais Dieu voit tout. Et le marché aussi a une mémoire.
  (Il sort. Kanga recompte ses billets.)

  Résultat final : scène de 15 répliques · 5 contraintes respectées
  · 4 didascalies : décor d'ouverture, ton (froidement), geste (levant les yeux),
    état (les mains tremblantes) · Tirade de Yéo (5 répliques) · Aparté final
  · Tonalité dramatique dominante · Contexte ivoirien (Bouaké, vivriers)
  · Note estimée : 17/20


✔ CORRIGÉ — EXERCICE 1
  Étape 1 · Didascalies (tout ce qui est entre parenthèses) :
    1 · (Une rue d'Abidjan. Il fait chaud. Soro et Diallo se croisent devant une boutique.)
        → didascalie de décor et de situation
    2 · (surpris) → didascalie d'état / de ton pour Soro
    3 · (froidement) → didascalie de ton pour Diallo
    4 · (déconcerté) → didascalie d'état pour Soro
    5 · (se retournant pour partir) → didascalie de mouvement pour Diallo
    6 · (Il sort. Soro reste seul, immobile.) → didascalie de sortie et d'état final
    Total : 6 didascalies

  Étape 2 · Répliques : 4 répliques
    (Soro / Diallo / Soro / Diallo)

  Étape 3 · Acte II · Scène 1

  Étape 4 · Sentiment dominant : tension / froideur / mystère
    Justification : « froidement » + « Trop longtemps. Ou pas assez. » (réplique énigmatique)
    + « Tu le sais très bien. » (sous-entendu non expliqué)

  Résultat final :
  A · 6 didascalies (décor · surprise · froideur · déconcertement · mouvement · sortie)
  B · 4 répliques
  C · Acte II · Scène 1
  D · Tension et mystère — justifiés par « froidement » et la réplique « Tu le sais très bien »


✔ CORRIGÉ — EXERCICE 2
  Étape 1 ·
  Situation A → MONOLOGUE
  Justification : « Seul sur scène » dans la didascalie · aucun interlocuteur présent
  Le personnage se parle à lui-même, révèle sa culpabilité

  Situation B → APARTÉ
  Justification : « à voix basse, tournée vers le public » + « pendant que les autres parlent »
  Les autres personnages sont présents mais ne sont pas censés entendre

  Situation C → TIRADE
  Justification : longue réplique d'un seul personnage (7 lignes) sans interruption
  Elle argumente et annonce un changement de situation

  Situation D → DIALOGUE
  Justification : échange rapide entre deux personnages (Assa et Moussa)
  Répliques courtes alternées — forme classique du dialogue

  Résultat final : A = monologue · B = aparté · C = tirade · D = dialogue


✔ CORRIGÉ — EXERCICE 3
  Étape 1 ·
  Situation 1 → Conflit EXTÉRIEUR (entre personnages)
  Opposition entre le père (choix économique) et la fille (choix sentimental)
  Deux personnes s'affrontent avec des volontés contraires dans chaque scène.

  Situation 2 → Conflit INTÉRIEUR (dans un personnage)
  Le fonctionnaire est seul face à un dilemme moral : dénoncer ou se taire.
  Le conflit se passe dans sa conscience, pas entre lui et une autre personne.

  Situation 3 → Conflit SOCIAL / POLITIQUE (entre groupes)
  Opposition entre les ouvriers (dominés) et le patron (dominant).
  Le conflit dépasse les individus : il met en jeu un rapport de force collectif
  et un système d'exploitation.

  Résultat final :
  Situation 1 = conflit extérieur · Situation 2 = conflit intérieur · Situation 3 = conflit social


✔ CORRIGÉ — EXERCICE 4
  Étape 1 ·
  A · Extrait A → Tonalité TRAGIQUE / PATHÉTIQUE
  Relevé 1 : « Mon fils… ils l'ont emmené cette nuit » — perte brutale et irrémédiable
  Relevé 2 : « Le ciel voit tout… mais il ne fait rien » — sentiment d'abandon et de fatalité
  La tonalité appelle à la compassion du spectateur face à une douleur insurmontable.

  B · Extrait B → Tonalité COMIQUE
  Type de comique : comique de situation (l'étiquette oubliée dans le dos)
  + comique de caractère (la vanité ridicule de M. Kassi qui se prend pour un président)
  Justification : « bombant le torse devant un miroir » + « l'étiquette qui pendait dans son dos »

  C ·
  L'Extrait A provoque de la tristesse et de la compassion — le théâtre devient ici
  un miroir de la souffrance humaine.
  L'Extrait B provoque le rire — le théâtre devient un espace de légèreté et de critique
  douce par l'absurde.
  Ces deux extraits montrent que le théâtre peut tout traiter : le deuil comme la vanité,
  le grave comme le ridicule. C'est cette diversité qui en fait un art universel.

  Résultat final :
  A · Tragique / pathétique — « ils l'ont emmené cette nuit » + « il ne fait rien »
  B · Comique — de situation (étiquette) + de caractère (vanité de Kassi)
  C · Le théâtre peut susciter la tristesse ou le rire selon ses choix stylistiques


✔ CORRIGÉ — EXERCICE 5
  Étape 1 ·
  A · 4 didascalies identifiées et classées :
  1 · (La case du chef. Soir. Une lampe éclaire faiblement…) → didascalie de DÉCOR
  2 · (sans le regarder) → didascalie de GESTE / d'attitude (mépris de Gnandé)
  3 · (avec un sourire) / (avec fausse compassion) → didascalies de TON
  4 · (serrant son chapeau, contenant sa colère) → didascalie d'état physique / émotion de Boua
  5 · (se levant, mettant fin à l'entretien) → didascalie de MOUVEMENT et d'intention
  6 · (dans l'embrasure, à voix basse, pour lui-même) → didascalie d'aparté

  B · La dernière réplique de Boua est un APARTÉ
  Justification : didascalie « à voix basse, pour lui-même » + il est dans l'embrasure
  (entre la scène et la sortie) — les autres personnages ne sont pas censés l'entendre

  C · Conflit dramatique :
  Boua (paysan, fils d'un cultivateur décédé) s'oppose à Gnandé (chef, détenteur du pouvoir)
  pour la possession d'un terrain cultivé par sa famille depuis trente ans.
  Conflit EXTÉRIEUR entre un homme du peuple et l'autorité locale.
  Conflit SOCIAL : opposition entre la justice coutumière (héritage du père) et l'arbitraire
  du pouvoir (les morts n'ont pas de terrain).

  D · Deux tonalités :
  Tonalité DRAMATIQUE : la situation est grave — Boua risque de perdre le moyen de nourrir
  ses enfants. Relevé : « ce terrain nourrit mes six enfants »
  Tonalité SATIRIQUE : Gnandé est ridiculisé par son cynisme excessif.
  Relevé : « Mes condoléances. Mais les morts n'ont pas de terrain. (Il rit.) »
  — La fausse compassion suivie du rire crée une critique mordante du pouvoir abusif.

  E · La didascalie « serrant son chapeau, contenant sa colère » est essentielle :
  Elle révèle ce que Boua ne peut pas dire à voix haute (sa rage, son humiliation).
  Le chapeau serré est un geste de retenue — il montre un homme qui se contrôle
  pour ne pas perdre le peu de dignité qui lui reste face au pouvoir.
  Sans cette didascalie, le spectateur ne verrait qu'un homme passif.
  Avec elle, il voit un homme qui souffre et résiste en silence.

  Résultat final :
  A · 6 didascalies : décor, geste de mépris, ton ironique, émotion contenue, mouvement final, aparté
  B · Aparté — confirmé par « à voix basse, pour lui-même »
  C · Conflit extérieur (Boua vs Gnandé) + social (peuple vs pouvoir arbitraire)
  D · Dramatique (« six enfants ») + Satirique (« les morts n'ont pas de terrain. Il rit. »)
  E · Le geste révèle la colère intérieure que la parole ne peut pas exprimer


════════════════════════════════════════════════════════
  CITATION FINALE
════════════════════════════════════════════════════════

  « Le théâtre, c'est la vie qui se regarde elle-même. »
  (Bernard Binlin Dadié)

  Prof. Mariam KONATÉ · Français · Seconde · DPFC/MENET-FP CI · 2025-2026

════════════════════════════════════════════════════════
"""

# ============================================================
#  GÉNÉRATEUR — VERSION 4 (style modèle de référence v2)
# ============================================================

import re
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── COULEURS ─────────────────────────────────────────────────
C_BLEU_TITRE  = RGBColor(0x18, 0x40, 0x82)
C_OR_TITRE    = RGBColor(0xC9, 0x92, 0x25)
C_GRIS_TEXTE  = RGBColor(0x23, 0x23, 0x23)
C_GRIS_CLAIR  = RGBColor(0x55, 0x55, 0x55)
C_VERT_OK     = RGBColor(0x2A, 0x7E, 0x5B)
C_ROUGE_ERR   = RGBColor(0xA6, 0x2D, 0x2D)
C_ORANGE_ANAL = RGBColor(0xC0, 0x54, 0x1A)
C_BLANC       = RGBColor(0xFF, 0xFF, 0xFF)

# ── FONDS ────────────────────────────────────────────────────
BG_SECTION    = RGBColor(0x18, 0x3E, 0x82)
BG_CAPSULE    = RGBColor(0xF0, 0xEB, 0xF8)
BG_ANCRAGE    = RGBColor(0xE2, 0xF0, 0xD9)
BG_HARKNESS   = RGBColor(0xEA, 0xF2, 0xF8)
BG_DECOUVERTE = RGBColor(0xEA, 0xF2, 0xF8)
BG_RETENIR    = RGBColor(0xE2, 0xF0, 0xD9)
BG_ATTENTION  = RGBColor(0xFF, 0xF2, 0xCC)
BG_DEF        = RGBColor(0xEA, 0xF2, 0xF8)
BG_TRACE      = RGBColor(0xED, 0xF4, 0xFB)
BG_ANALOGIE   = RGBColor(0xFE, 0xF0, 0xE7)
BG_SYNTHESE   = RGBColor(0xD9, 0xE2, 0xF3)
BG_EXO_GUIDE  = RGBColor(0xF5, 0xF7, 0xFA)
BG_EXO_ENT    = RGBColor(0xFA, 0xFA, 0xFA)
BG_DEVOIR     = RGBColor(0xFF, 0xF9, 0xEE)
BG_CORRIGE    = RGBColor(0xF8, 0xFB, 0xF7)
BG_SCHEMA     = RGBColor(0xFA, 0xFA, 0xFA)
BG_SOMMAIRE   = RGBColor(0xEE, 0xF3, 0xFA)


def _hex(c):
    return f"{c[0]:02X}{c[1]:02X}{c[2]:02X}"


def set_cell_bg(cell, color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for e in tcPr.findall(qn('w:shd')):
        tcPr.remove(e)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), _hex(color))
    tcPr.append(shd)


def remove_cell_borders(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "none")
        b.set(qn("w:sz"), "0")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "auto")
        tcBorders.append(b)
    tcPr.append(tcBorders)


def add_cell_border_bottom(cell, color="184082", sz="4"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), sz)
    bottom.set(qn("w:space"), "0")
    bottom.set(qn("w:color"), color)
    tcBorders.append(bottom)
    tcPr.append(tcBorders)


def add_cell_border_top(cell, color="184082", sz="4"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), sz)
    top.set(qn("w:space"), "0")
    top.set(qn("w:color"), color)
    tcBorders.append(top)
    tcPr.append(tcBorders)


doc = Document()
for sec in doc.sections:
    sec.top_margin    = Cm(1.8)
    sec.bottom_margin = Cm(1.8)
    sec.left_margin   = Cm(2.0)
    sec.right_margin  = Cm(2.0)


# ── EN-TÊTE ──────────────────────────────────────────────────
def add_header(lecon_info):
    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False
    tbl = header.add_table(rows=1, cols=3, width=Cm(17))
    tbl.style = "Table Grid"
    cells = tbl.rows[0].cells
    for cell in cells:
        remove_cell_borders(cell)
        add_cell_border_bottom(cell)
    p0 = cells[0].paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r0 = p0.add_run("LPA · 3ème")
    r0.font.size = Pt(8); r0.font.color.rgb = C_GRIS_CLAIR
    p1 = cells[1].paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run(lecon_info)
    r1.bold = True; r1.font.size = Pt(8); r1.font.color.rgb = C_BLEU_TITRE
    p2 = cells[2].paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r2 = p2.add_run("Anglais LV1")
    r2.font.size = Pt(8); r2.font.color.rgb = C_GRIS_CLAIR


# ── PIED DE PAGE ─────────────────────────────────────────────
def add_footer(prof, organisme):
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    tbl = footer.add_table(rows=1, cols=3, width=Cm(17))
    tbl.style = "Table Grid"
    cells = tbl.rows[0].cells
    for cell in cells:
        remove_cell_borders(cell)
        add_cell_border_top(cell)
    p0 = cells[0].paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r0 = p0.add_run(prof); r0.font.size = Pt(8); r0.font.color.rgb = C_GRIS_CLAIR
    p1 = cells[1].paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run(organisme); r1.font.size = Pt(8); r1.font.color.rgb = C_GRIS_CLAIR
    p2 = cells[2].paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r2 = p2.add_run("Page "); r2.font.size = Pt(8); r2.font.color.rgb = C_GRIS_CLAIR
    fldChar1 = OxmlElement("w:fldChar"); fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText"); instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar"); fldChar2.set(qn("w:fldCharType"), "end")
    r_p = p2.add_run(); r_p.font.size = Pt(8); r_p.font.color.rgb = C_GRIS_CLAIR
    r_p._r.append(fldChar1); r_p._r.append(instrText); r_p._r.append(fldChar2)


# ── PAGE DE GARDE ────────────────────────────────────────────
def add_page_garde(titre, competence, professeur, organisme, lecon_num):
    # Grand LPA
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(50)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run("LPA")
    r.bold = True; r.font.size = Pt(60); r.font.color.rgb = C_BLEU_TITRE
    # Sous-titre
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after  = Pt(16)
    r2 = p2.add_run("Lycée Personnel Autonome")
    r2.italic = True; r2.font.size = Pt(13); r2.font.color.rgb = C_GRIS_CLAIR
    # Ligne bleue
    p_sep = doc.add_paragraph()
    p_sep.paragraph_format.space_before = Pt(0)
    p_sep.paragraph_format.space_after  = Pt(14)
    pPr = p_sep._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '16')
    bottom.set(qn('w:space'), '1'); bottom.set(qn('w:color'), '184082')
    pBdr.append(bottom); pPr.append(pBdr)
    # Matière · Leçon · Niveau
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_before = Pt(0)
    p3.paragraph_format.space_after  = Pt(6)
    r3 = p3.add_run(f"Anglais LV1  ·  {lecon_num}  ·  Classe de 3ème")
    r3.bold = True; r3.font.size = Pt(14); r3.font.color.rgb = C_BLEU_TITRE
    # Grand titre — toujours affiché, même si vide on prend l'en-tête
    titre_affiche = titre if titre else lecon_num
    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p4.paragraph_format.space_before = Pt(0)
    p4.paragraph_format.space_after  = Pt(6)
    r4 = p4.add_run(titre_affiche)
    r4.bold = True; r4.font.size = Pt(32); r4.font.color.rgb = C_OR_TITRE
    # Compétence italique
    p5 = doc.add_paragraph()
    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p5.paragraph_format.space_before = Pt(0)
    p5.paragraph_format.space_after  = Pt(30)
    r5 = p5.add_run(competence)
    r5.italic = True; r5.font.size = Pt(12); r5.font.color.rgb = C_GRIS_CLAIR
    # Professeur
    p6 = doc.add_paragraph()
    p6.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p6.paragraph_format.space_before = Pt(0)
    p6.paragraph_format.space_after  = Pt(6)
    r6 = p6.add_run(professeur)
    r6.font.size = Pt(12); r6.font.color.rgb = C_GRIS_TEXTE
    # Organisme
    p7 = doc.add_paragraph()
    p7.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p7.paragraph_format.space_before = Pt(0)
    p7.paragraph_format.space_after  = Pt(6)
    r7 = p7.add_run(organisme)
    r7.font.size = Pt(11); r7.font.color.rgb = C_GRIS_CLAIR
    # Saut de page
    doc.add_page_break()


# ── SOMMAIRE ─────────────────────────────────────────────────
def add_sommaire(sommaire_lines, sections_phases):
    """
    sommaire_lines : lignes du cours (1. Avant la leçon, etc.)
    sections_phases : dict pour enrichir avec les phases
    """
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, BG_SOMMAIRE)
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in [("top","80"),("bottom","80"),("left","140"),("right","140")]:
        m = OxmlElement(f"w:{side}")
        m.set(qn("w:w"), val); m.set(qn("w:type"), "dxa")
        tcMar.append(m)
    tcPr.append(tcMar)
    # Titre Sommaire
    p0 = cell.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p0.paragraph_format.space_before = Pt(0)
    p0.paragraph_format.space_after  = Pt(8)
    r0 = p0.add_run("Sommaire")
    r0.bold = True; r0.font.size = Pt(16); r0.font.color.rgb = C_BLEU_TITRE

    # Mapping sections → sous-points enrichis
    enriched = {
        "1": ["· Capsule de reprise", "· Ancrage ivoirien", "· Questions Harkness"],
        "2": ["· Phase 1 — Présentation & prérequis", "· Phase 2 — Découverte guidée",
              "· Phase 3 — Schémas textuels", "· Phase 4 — Définitions DPFC",
              "· Phase 5 — Trace écrite + Analogies CI", "· Phase 5b — Synthèse",
              "· Phase 6 — Exercices guidés"],
        "3": ["· Exercices 1 à 5 + Devoir maison"],
        "4": ["· Corrigé du devoir maison + Exercices 1 à 5"],
    }

    num_auto = 0  # compteur si pas de numéro explicite
    for line in sommaire_lines:
        line = line.strip()
        if not line:
            continue
        # Nouveau format : "Section N — ..."
        m2 = re.match(r'^Section\s+(\d+)\s*[—–-]\s*(.*)', line, re.IGNORECASE)
        if m2:
            num = m2.group(1); label = m2.group(2).strip()
            p = cell.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(2)
            r = p.add_run(f"Section {num} — {label}")
            r.bold = True; r.font.size = Pt(11); r.font.color.rgb = C_BLEU_TITRE
            for sub in enriched.get(num, []):
                ps = cell.add_paragraph()
                ps.paragraph_format.space_before = Pt(0)
                ps.paragraph_format.space_after  = Pt(1)
                ps.paragraph_format.left_indent  = Pt(14)
                rs = ps.add_run(sub)
                rs.font.size = Pt(10); rs.font.color.rgb = C_GRIS_CLAIR
            continue
        # Sous-point du nouveau format : "    · Phase..."
        if re.match(r'^[·•]\s+', line):
            ps = cell.add_paragraph()
            ps.paragraph_format.space_before = Pt(0)
            ps.paragraph_format.space_after  = Pt(1)
            ps.paragraph_format.left_indent  = Pt(14)
            rs = ps.add_run(line)
            rs.font.size = Pt(10); rs.font.color.rgb = C_GRIS_CLAIR
            continue
        m = re.match(r'^(\d+)\.\s*(.*)', line)
        if m:
            num = m.group(1); label = m.group(2)
            # Titre de section
            p = cell.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(2)
            r = p.add_run(f"Section {num} — {label}")
            r.bold = True; r.font.size = Pt(11); r.font.color.rgb = C_BLEU_TITRE
            # Sous-points
            for sub in enriched.get(num, []):
                ps = cell.add_paragraph()
                ps.paragraph_format.space_before = Pt(0)
                ps.paragraph_format.space_after  = Pt(1)
                ps.paragraph_format.left_indent  = Pt(14)
                rs = ps.add_run(sub)
                rs.font.size = Pt(10); rs.font.color.rgb = C_GRIS_CLAIR

    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    doc.add_page_break()


# ── HELPERS ───────────────────────────────────────────────────
def _run(para, text, bold=False, italic=False, color=None, size=10.5):
    r = para.add_run(text)
    r.bold = bold; r.italic = italic
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    return r


def make_box(bg_color):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, bg_color)
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in [("top","70"),("bottom","70"),("left","120"),("right","120")]:
        m = OxmlElement(f"w:{side}")
        m.set(qn("w:w"), val); m.set(qn("w:type"), "dxa")
        tcMar.append(m)
    tcPr.append(tcMar)
    return tbl, cell


def space_after_box(pts=4):
    doc.add_paragraph().paragraph_format.space_after = Pt(pts)


def _add_para_in_cell(cell, text, bold=False, italic=False,
                       color=None, size=10.5, sb=0, sa=2,
                       align=WD_ALIGN_PARAGRAPH.LEFT, indent_pt=0):
    p = cell.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(sa)
    if indent_pt:
        p.paragraph_format.left_indent = Pt(indent_pt)
    _run(p, text, bold=bold, italic=italic,
         color=color or C_GRIS_TEXTE, size=size)
    return p


# ── BANNIÈRE SECTION ─────────────────────────────────────────
def add_banner(text):
    tbl, cell = make_box(BG_SECTION)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    _run(p, text, bold=True, color=C_BLANC, size=13)
    space_after_box(5)


# ── PHASE HEADER ─────────────────────────────────────────────
def add_phase(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    # Ligne de soulignement
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '2')
    bottom.set(qn('w:space'), '1'); bottom.set(qn('w:color'), 'D0D8E8')
    pBdr.append(bottom); pPr.append(pBdr)
    _run(p, text, bold=True, color=C_BLEU_TITRE, size=14)


# ── SOUS-TITRE ───────────────────────────────────────────────
def add_sous_titre(text, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after  = Pt(2)
    _run(p, text, bold=True, color=color or C_OR_TITRE, size=11.5)


# ── TEXTE COURANT ────────────────────────────────────────────
def add_text(text, italic=False, color=None, size=11, sb=0, sa=3,
             align=WD_ALIGN_PARAGRAPH.LEFT, indent_pt=0):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(sa)
    if indent_pt:
        p.paragraph_format.left_indent = Pt(indent_pt)
    _run(p, text, italic=italic, color=color or C_GRIS_TEXTE, size=size)


# ── BOÎTE GÉNÉRIQUE ──────────────────────────────────────────
def add_simple_box(lines, bg, title=None, title_color=None,
                   title_size=10, title_icon=""):
    tbl, cell = make_box(bg)
    first = True
    if title:
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(5)
        txt = (title_icon + "  " + title) if title_icon else title
        _run(p, txt, bold=True, color=title_color or C_BLEU_TITRE, size=title_size)
        first = False
    for line in lines:
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(2)
        _run(p, line, color=C_GRIS_TEXTE, size=10.5)
    space_after_box()


# ── BOÎTE HARKNESS enrichie ──────────────────────────────────
def add_harkness_box(title, lines):
    tbl, cell = make_box(BG_HARKNESS)
    # Titre
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(5)
    _run(p, "▶  " + title, bold=True, color=C_BLEU_TITRE, size=12)
    for line in lines:
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(2)
        if line.startswith("Q :") or line.startswith("Q:"):
            _run(p, line, color=C_GRIS_TEXTE, size=10.5)
        elif line.startswith("Guide") or line.startswith("Piste"):
            _run(p, line, italic=True, color=C_GRIS_CLAIR, size=10.5)
        elif line.startswith("Corrigé") or line.startswith("Corrige"):
            _run(p, line, bold=True, color=C_VERT_OK, size=10.5)
        else:
            _run(p, line, color=C_GRIS_TEXTE, size=10.5)
    space_after_box()


# ── BOÎTE TEXTE DE DÉCOUVERTE ────────────────────────────────
def add_decouverte_box(eng_lines, trad_lines):
    tbl, cell = make_box(BG_DECOUVERTE)
    p0 = cell.paragraphs[0]
    p0.paragraph_format.space_before = Pt(0)
    p0.paragraph_format.space_after  = Pt(5)
    _run(p0, "📖  Texte de découverte", bold=True, color=C_BLEU_TITRE, size=10)
    for line in eng_lines:
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(2)
        _run(p, line, color=C_GRIS_TEXTE, size=10.5)
    if trad_lines:
        # Ligne séparatrice visuelle
        p_sep = cell.add_paragraph()
        p_sep.paragraph_format.space_before = Pt(6)
        p_sep.paragraph_format.space_after  = Pt(4)
        pPr = p_sep._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        top_b = OxmlElement('w:top')
        top_b.set(qn('w:val'), 'single'); top_b.set(qn('w:sz'), '2')
        top_b.set(qn('w:space'), '1'); top_b.set(qn('w:color'), 'B0C4D8')
        pBdr.append(top_b); pPr.append(pBdr)
        _run(p_sep, "— Traduction :", bold=True, color=C_BLEU_TITRE, size=10)
        for line in trad_lines:
            p = cell.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(2)
            _run(p, line, italic=True, color=C_GRIS_CLAIR, size=10.5)
    space_after_box()


# ── BOÎTE À RETENIR ──────────────────────────────────────────
def add_retenir_box(title, lines):
    tbl, cell = make_box(BG_RETENIR)
    p0 = cell.paragraphs[0]
    p0.paragraph_format.space_before = Pt(0)
    p0.paragraph_format.space_after  = Pt(5)
    _run(p0, "📘  " + title, bold=True, color=C_VERT_OK, size=10)
    for line in lines:
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(2)
        _run(p, line, color=C_GRIS_TEXTE, size=10.5)
    space_after_box()


# ── BOÎTE ATTENTION ──────────────────────────────────────────
def add_attention_box(title, lines):
    tbl, cell = make_box(BG_ATTENTION)
    p0 = cell.paragraphs[0]
    p0.paragraph_format.space_before = Pt(0)
    p0.paragraph_format.space_after  = Pt(5)
    _run(p0, "⚠  " + title, bold=True, color=C_OR_TITRE, size=10)
    for line in lines:
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(2)
        # ✗ en rouge, ✓ en vert
        if '✗' in line or '✓' in line:
            parts = re.split(r'(→)', line)
            for part in parts:
                part = part.strip()
                if not part:
                    continue
                if '✗' in part:
                    _run(p, part + "  ", color=C_ROUGE_ERR, size=10.5)
                elif '✓' in part:
                    _run(p, part, bold=True, color=C_VERT_OK, size=10.5)
                elif part == '→':
                    _run(p, "  →  ", italic=True, color=C_GRIS_CLAIR, size=10.5)
                else:
                    _run(p, part, color=C_GRIS_TEXTE, size=10.5)
        else:
            _run(p, line, color=C_GRIS_TEXTE, size=10.5)
    space_after_box()


# ── BOÎTE DÉFINITION ─────────────────────────────────────────
def add_def_box(title, lines):
    tbl, cell = make_box(BG_DEF)
    p0 = cell.paragraphs[0]
    p0.paragraph_format.space_before = Pt(0)
    p0.paragraph_format.space_after  = Pt(5)
    _run(p0, title, bold=True, color=C_BLEU_TITRE, size=11)
    for line in lines:
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(2)
        _run(p, line, color=C_GRIS_TEXTE, size=10.5)
    space_after_box()


# ── BOÎTE TRACE ÉCRITE ───────────────────────────────────────
def add_trace_box(lines):
    tbl, cell = make_box(BG_TRACE)
    first = True
    for line in lines:
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(2)
        if re.match(r'^(I|II|III|IV|V)\.\s', line):
            _run(p, line, bold=True, color=C_BLEU_TITRE, size=11)
        elif re.match(r'^✎', line):
            _run(p, line, bold=True, color=C_BLEU_TITRE, size=10)
        elif re.match(r'^⚠', line):
            _run(p, line, bold=True, color=C_OR_TITRE, size=10.5)
        else:
            _run(p, line, color=C_GRIS_TEXTE, size=10.5)
    space_after_box()


# ── BOÎTE ANALOGIE ───────────────────────────────────────────
def add_analogie_box(title, lines):
    tbl, cell = make_box(BG_ANALOGIE)
    p0 = cell.paragraphs[0]
    p0.paragraph_format.space_before = Pt(0)
    p0.paragraph_format.space_after  = Pt(5)
    _run(p0, "💡  " + title, bold=True, color=C_ORANGE_ANAL, size=10)
    for line in lines:
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(2)
        _run(p, line, color=C_GRIS_TEXTE, size=10.5)
    space_after_box()


# ── BOÎTE SYNTHÈSE ───────────────────────────────────────────
def add_synthese_box(lines):
    tbl, cell = make_box(BG_SYNTHESE)
    first = True
    for line in lines:
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(2)
        if re.match(r'^5 points|^5 mots|^3 erreurs', line):
            _run(p, line, bold=True, color=C_BLEU_TITRE, size=11)
        elif re.match(r'^\s*\d+\.', line):
            m = re.match(r'^(\s*\d+\.\s*)(.*)', line)
            if m:
                _run(p, m.group(1), bold=True, color=C_BLEU_TITRE, size=10.5)
                _run(p, m.group(2), color=C_GRIS_TEXTE, size=10.5)
        elif re.match(r'^\s*·', line):
            m = re.match(r'^(\s*·\s*)(\S+)(.*)', line)
            if m:
                _run(p, m.group(1), bold=True, color=C_OR_TITRE, size=10.5)
                _run(p, m.group(2), bold=True, color=C_GRIS_TEXTE, size=10.5)
                _run(p, m.group(3), color=C_GRIS_CLAIR, size=10.5)
        elif re.match(r'^\s*✗', line):
            m = re.match(r'^(\s*✗\s*)(.*?)\s*→\s*(✓.*)', line)
            if m:
                _run(p, m.group(1) + m.group(2) + "  ", bold=True, color=C_ROUGE_ERR, size=10.5)
                _run(p, "→  ", italic=True, color=C_GRIS_CLAIR, size=10.5)
                _run(p, m.group(3), bold=True, color=C_VERT_OK, size=10.5)
            else:
                _run(p, line, color=C_ROUGE_ERR, size=10.5)
        else:
            _run(p, line, color=C_GRIS_TEXTE, size=10.5)
    space_after_box()


# ── BOÎTE EXERCICE GUIDÉ ─────────────────────────────────────
def add_exo_guide_box(title, lines):
    tbl, cell = make_box(BG_EXO_GUIDE)
    p0 = cell.paragraphs[0]
    p0.paragraph_format.space_before = Pt(0)
    p0.paragraph_format.space_after  = Pt(5)
    _run(p0, "◉  " + title, bold=True, color=C_BLEU_TITRE, size=12)
    for line in lines:
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(2)
        if re.match(r'^Notion ciblée\s*:', line):
            m = re.match(r'^(Notion ciblée\s*:\s*)(.*)', line)
            if m:
                _run(p, m.group(1), bold=True, color=C_OR_TITRE, size=10.5)
                _run(p, m.group(2), italic=True, color=C_GRIS_TEXTE, size=10.5)
        elif re.match(r'^(Énoncé|MÉTHODE|Démarche)', line):
            _run(p, line, bold=True, color=C_GRIS_TEXTE, size=10.5)
        elif re.match(r'^Étape \d+', line):
            m = re.match(r'^(Étape \d+ · )(.*)', line)
            if m:
                _run(p, m.group(1), bold=True, color=C_BLEU_TITRE, size=10.5)
                _run(p, m.group(2), color=C_GRIS_TEXTE, size=10.5)
            else:
                _run(p, line, color=C_GRIS_TEXTE, size=10.5)
        elif re.match(r'^✔', line):
            m = re.match(r'^(✔[^:]*:\s*)(.*)', line)
            if m:
                _run(p, m.group(1), bold=True, color=C_VERT_OK, size=10.5)
                _run(p, m.group(2), color=C_GRIS_TEXTE, size=10.5)
            else:
                _run(p, line, bold=True, color=C_VERT_OK, size=10.5)
        else:
            _run(p, line, color=C_GRIS_TEXTE, size=10.5)
    space_after_box()


# ── BOÎTE EXERCICE ENTRAÎNEMENT ──────────────────────────────
def add_exo_ent_box(title, lines):
    tbl, cell = make_box(BG_EXO_ENT)
    p0 = cell.paragraphs[0]
    p0.paragraph_format.space_before = Pt(0)
    p0.paragraph_format.space_after  = Pt(5)
    _run(p0, "◎  " + title, bold=True, color=C_BLEU_TITRE, size=12)
    for line in lines:
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(2)
        if re.match(r'^Notions travaillées\s*:', line):
            m = re.match(r'^(Notions travaillées\s*:\s*)(.*)', line)
            if m:
                _run(p, m.group(1), bold=True, color=C_OR_TITRE, size=10.5)
                _run(p, m.group(2), italic=True, color=C_GRIS_TEXTE, size=10.5)
        elif re.match(r'^GUIDE|^Guide de réussite', line):
            _run(p, line, bold=True, color=C_OR_TITRE, size=10.5)
        elif re.match(r'^Étape \d+', line):
            m = re.match(r'^(Étape \d+ · )(.*)', line)
            if m:
                _run(p, m.group(1), bold=True, color=C_BLEU_TITRE, size=10.5)
                _run(p, m.group(2), color=C_GRIS_TEXTE, size=10.5)
            else:
                _run(p, line, color=C_GRIS_TEXTE, size=10.5)
        else:
            _run(p, line, color=C_GRIS_TEXTE, size=10.5)
    space_after_box()


# ── BOÎTE DEVOIR MAISON ──────────────────────────────────────
def add_devoir_box(title, lines):
    tbl, cell = make_box(BG_DEVOIR)
    p0 = cell.paragraphs[0]
    p0.paragraph_format.space_before = Pt(0)
    p0.paragraph_format.space_after  = Pt(5)
    _run(p0, "◈  " + title, bold=True, color=C_OR_TITRE, size=12)
    for line in lines:
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(2)
        if re.match(r'^(CONTEXTE|ÉNONCÉ|ENONCE|STRUCTURE|CONTRAINTES|GUIDE PAS)', line):
            _run(p, line, bold=True, color=C_OR_TITRE, size=10.5)
        elif re.match(r'^Durée', line):
            _run(p, line, italic=True, color=C_GRIS_CLAIR, size=10)
        elif re.match(r'^Étape \d+', line):
            m = re.match(r'^(Étape \d+ · )(.*)', line)
            if m:
                _run(p, m.group(1), bold=True, color=C_OR_TITRE, size=10.5)
                _run(p, m.group(2), color=C_GRIS_TEXTE, size=10.5)
            else:
                _run(p, line, color=C_GRIS_TEXTE, size=10.5)
        elif line.startswith('·'):
            _run(p, line, color=C_GRIS_TEXTE, size=10.5)
        elif re.match(r'^\d+\.', line):
            _run(p, line, color=C_GRIS_TEXTE, size=10.5)
        else:
            _run(p, line, color=C_GRIS_TEXTE, size=10.5)
    space_after_box()


# ── BOÎTE CORRIGÉ ────────────────────────────────────────────
def add_corrige_box(title, lines):
    tbl, cell = make_box(BG_CORRIGE)
    p0 = cell.paragraphs[0]
    p0.paragraph_format.space_before = Pt(0)
    p0.paragraph_format.space_after  = Pt(5)
    _run(p0, "✔  " + title, bold=True, color=C_VERT_OK, size=12)
    for line in lines:
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(2)
        if re.match(r'^✔', line):
            _run(p, line, bold=True, color=C_VERT_OK, size=10.5)
        elif re.match(r'^TRADUCTION|^TITRE PROPOSÉ|^TEXTE PROPOSÉ', line):
            _run(p, line, bold=True, color=C_GRIS_TEXTE, size=10.5)
        else:
            _run(p, line, color=C_GRIS_TEXTE, size=10.5)
    space_after_box()


# ── TABLEAU STRUCTURES FONDAMENTALES ────────────────────────
def add_tableau_box(title, md_lines):
    """Transforme les lignes markdown | col | col | en vrai tableau Word."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn as _qn

    # Parser les lignes markdown en liste de colonnes
    rows_data = []
    for line in md_lines:
        line = line.strip().strip('|')
        cols = [c.strip() for c in line.split('|')]
        if cols:
            rows_data.append(cols)

    if not rows_data:
        return

    # Titre optionnel au-dessus
    if title:
        p_t = doc.add_paragraph()
        p_t.paragraph_format.space_before = Pt(6)
        p_t.paragraph_format.space_after  = Pt(3)
        _run(p_t, title, bold=True, color=C_BLEU_TITRE, size=11)

    nb_cols = max(len(r) for r in rows_data)
    tbl = doc.add_table(rows=len(rows_data), cols=nb_cols)
    tbl.style = "Table Grid"

    for r_idx, row_data in enumerate(rows_data):
        is_header = (r_idx == 0)
        row = tbl.rows[r_idx]
        for c_idx in range(nb_cols):
            cell = row.cells[c_idx]
            text = row_data[c_idx] if c_idx < len(row_data) else ""
            # Fond
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(_qn("w:val"), "clear")
            shd.set(_qn("w:color"), "auto")
            shd.set(_qn("w:fill"), "183E82" if is_header else ("F0F4FA" if r_idx % 2 == 0 else "FFFFFF"))
            tcPr.append(shd)
            # Texte
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            r = p.add_run(text)
            r.bold = is_header
            r.font.size = Pt(9.5)
            r.font.color.rgb = C_BLANC if is_header else C_GRIS_TEXTE

    doc.add_paragraph().paragraph_format.space_after = Pt(5)


# ── CITATION FINALE ──────────────────────────────────────────
def add_citation(text_en, text_fr, prof_org):
    doc.add_paragraph().paragraph_format.space_after = Pt(10)
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_before = Pt(4)
    p1.paragraph_format.space_after  = Pt(2)
    _run(p1, f'"{text_en}"', italic=True, color=C_BLEU_TITRE, size=11)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after  = Pt(2)
    _run(p2, f'({text_fr})', italic=True, color=C_GRIS_CLAIR, size=10)
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_before = Pt(2)
    p3.paragraph_format.space_after  = Pt(0)
    _run(p3, prof_org, color=C_GRIS_CLAIR, size=9)


# ============================================================
#  PARSEUR PRINCIPAL
# ============================================================

lines     = COURS.strip().split('\n')
total     = len(lines)
meta      = {}

# ── Extraction en-tête ════...════ ──────────────────────────
header_lines = []
in_header = False; header_done = False; body_start = 0
for idx, ln in enumerate(lines):
    s = ln.strip()
    if re.match(r'^[═]{4,}', s) and not in_header and not header_done:
        in_header = True; continue
    if re.match(r'^[═]{4,}', s) and in_header:
        in_header = False; header_done = True; body_start = idx + 1; break
    if in_header and s:
        header_lines.append(s)

# ── Méta-lignes ─────────────────────────────────────────────
i = body_start
while i < total:
    s = lines[i].strip()
    m = re.match(r'^(TITRE_LECON|COMPETENCE|PROFESSEUR|ORGANISME)\s*:\s*(.+)', s)
    if m:
        meta[m.group(1)] = m.group(2).strip(); i += 1
    elif not s or s == "SOMMAIRE":
        i += 1
        if s == "SOMMAIRE": break
    else:
        break

titre_lecon = meta.get("TITRE_LECON", "")
competence  = meta.get("COMPETENCE",  "")
professeur  = meta.get("PROFESSEUR",  "")
organisme   = meta.get("ORGANISME",   "")

# Si TITRE_LECON absent : extraire depuis le sommaire (ligne "Section 2")
# ou depuis l'en-tête ════ après le "—"
if not titre_lecon and header_lines:
    # Chercher "Leçon N — TITRE" dans la première ligne d'en-tête
    m_titre = re.search(r'[—–-]\s*(.+)$', header_lines[0])
    if m_titre:
        titre_lecon = m_titre.group(1).strip()

# Si PROFESSEUR absent : chercher dans l'en-tête
if not professeur and header_lines:
    for hl in header_lines:
        m_p = re.search(r'Professeur\s*:\s*(.+)', hl)
        if m_p:
            professeur = m_p.group(1).strip(); break

# Si COMPETENCE absent : chercher dans l'en-tête
if not competence and header_lines:
    for hl in header_lines:
        m_c = re.search(r'Compétence\s*:\s*(.+)', hl)
        if m_c:
            competence = m_c.group(1).strip(); break

# Si ORGANISME absent : chercher dans l'en-tête
if not organisme and header_lines:
    for hl in header_lines:
        m_o = re.search(r'(DPFC/.+)', hl)
        if m_o:
            organisme = m_o.group(1).strip(); break

# Numéro de leçon depuis l'en-tête
lecon_num = "Leçon"
if header_lines:
    m_lec = re.search(r'(Leçon\s*\d+)', header_lines[0])
    if m_lec:
        lecon_num = m_lec.group(1)

# Sommaire
sommaire_lines = []
while i < total:
    s = lines[i].strip()
    if not s or re.match(r'^[─═]{4,}', s): break
    sommaire_lines.append(s); i += 1

# ── En-tête / pied / page de garde / sommaire ───────────────
lecon_info = header_lines[0] if header_lines else "Anglais LV1"
add_header(lecon_info)
add_footer(professeur, organisme)
add_page_garde(titre_lecon, competence, professeur, organisme, lecon_num)
add_sommaire(sommaire_lines, {})

# ── Variables d'état ────────────────────────────────────────
cur_type  = None
cur_title = ""
cur_lines = []
in_tilde  = False
tilde_lines = []
in_syn    = False
syn_lines = []
in_trad_libre  = False
trad_libre_lines = []

SOUS_TITRES = re.compile(
    r'^(Objectifs|Prérequis|MÉTHODE|GUIDE|CONTRAINTES|CONTEXTE|'
    r'CONSIGNE|STRUCTURE|GUIDE PAS|Énoncé)')


def flush():
    global cur_type, cur_title, cur_lines
    if not cur_type:
        return
    t, ti, ls = cur_type, cur_title, cur_lines

    if   t == "CAPSULE":   add_simple_box(ls, BG_CAPSULE, title=ti,
                               title_color=C_BLEU_TITRE, title_size=10, title_icon="📌")
    elif t == "ANCRAGE":   add_simple_box(ls, BG_ANCRAGE, title=ti,
                               title_color=C_BLEU_TITRE, title_size=10, title_icon="🌍")
    elif t == "HARKNESS":  add_harkness_box(ti, ls)
    elif t == "RETENIR":   add_retenir_box(ti, ls)
    elif t == "ATTENTION": add_attention_box(ti, ls)
    elif t == "DEF":       add_def_box(ti, ls)
    elif t == "SCHEMA":    add_simple_box(ls, BG_SCHEMA, title=ti,
                               title_color=C_BLEU_TITRE, title_size=10, title_icon="📊")
    elif t == "TRACE":     add_trace_box(ls)
    elif t == "ANALOGIE":  add_analogie_box(ti, ls)
    elif t == "DEVOIR":    add_devoir_box(ti, ls)
    elif t == "EXOGUIDE":  add_exo_guide_box(ti, ls)
    elif t == "EXOENT":    add_exo_ent_box(ti, ls)
    elif t == "CORRIGE":   add_corrige_box(ti, ls)
    elif t == "TABLEAU":   add_tableau_box(ti, ls)

    cur_type = None; cur_title = ""; cur_lines = []


# ── Citation finale : stockage ───────────────────────────────
citation_en = "Speak even if it's not perfect — silence teaches nothing."
citation_fr = "Parle même si c'est imparfait — le silence n'apprend rien."
citation_sig = ""

# ── Boucle principale ────────────────────────────────────────
while i < total:
    raw = lines[i]; s = raw.strip(); i += 1
    if not s:
        continue

    # Méta déjà traitées
    if re.match(r'^(TITRE_LECON|COMPETENCE|PROFESSEUR|ORGANISME)\s*:', s):
        continue

    # Bloc ~~~ texte de découverte
    if s == "~~~":
        if not in_tilde:
            in_tilde = True; tilde_lines = []
        else:
            in_tilde = False
            flush()
            eng_lines = []; trad_lines = []; in_trad = False
            for tl in tilde_lines:
                if tl == "---TRADUCTION---":
                    in_trad = True; continue
                (trad_lines if in_trad else eng_lines).append(tl)
            # Si la traduction n'était pas dans les ~~~,
            # on active le mode "capture traduction libre"
            add_decouverte_box(eng_lines, trad_lines)
            if not trad_lines:
                in_trad_libre = True
                trad_libre_lines = []
        continue
    if in_tilde:
        tilde_lines.append(s); continue

    # Capture traduction libre après ~~~ sans ---TRADUCTION---
    if in_trad_libre:
        # On arrête à la prochaine boîte ou séparateur
        if (re.match(r'^[┌⚠◆◉◎◈✔✎~▶►\[]', s) or
                re.match(r'^(SECTION|Phase|SYNTHESE)', s) or
                re.match(r'^[─═]{4,}', s)):
            # Injecter les lignes de traduction dans la dernière boîte découverte
            # (on les ajoute via add_text en style traduit)
            in_trad_libre = False
            # Relancer le traitement de cette ligne (ne pas ignorer)
            # on laisse tomber dans la suite du parseur
        else:
            # Ignorer "TRADUCTION COMPLÈTE :" et absorber les lignes
            if not re.match(r'^TRADUCTION', s):
                trad_libre_lines.append(s)
            continue

    # SYNTHESE — avec balises explicites OU détection automatique
    if s == "SYNTHESE_DEBUT" or re.match(r'^5 points essentiels', s):
        flush(); in_syn = True; syn_lines = []
        if s != "SYNTHESE_DEBUT":
            syn_lines.append(s)  # inclure la ligne déclencheuse
        continue
    if s == "SYNTHESE_FIN":
        in_syn = False; add_synthese_box(syn_lines); syn_lines = []; continue
    if in_syn:
        # Fin automatique de synthèse à la prochaine Phase ou Section
        if re.match(r'^(Phase \d|SECTION)', s):
            in_syn = False; add_synthese_box(syn_lines); syn_lines = []
            # Ne pas ignorer cette ligne, la traiter ci-dessous
        else:
            syn_lines.append(s); continue

    # Séparateurs
    if re.match(r'^[═]{4,}', s) or re.match(r'^[─]{4,}', s):
        flush(); continue

    # SECTION
    if re.match(r'^SECTION \d+\s*[—–-]', s):
        flush(); add_banner(s); continue

    # Phase
    if re.match(r'^Phase \d+[a-z]?\s*[·.]', s):
        flush(); add_phase(s); continue

    # [BLOC TITRE] avec crochets
    m = re.match(r'^\[(.+)\]$', s)
    if m:
        flush(); bl = m.group(1); tu = bl.upper()
        if "CAPSULE" in tu:
            cur_type = "CAPSULE";  cur_title = bl
        elif "ANCRAGE" in tu:
            cur_type = "ANCRAGE";  cur_title = bl
        elif "SCHEMA" in tu or "SCHÉMA" in tu:
            cur_type = "SCHEMA";   cur_title = bl
        else:
            cur_type = "SCHEMA";   cur_title = bl
        continue

    # CAPSULE sans crochets : "CAPSULE DE REPRISE — ..."
    if re.match(r'^CAPSULE\s+DE\s+REPRISE', s):
        flush()
        cur_type = "CAPSULE"; cur_title = s.strip(); cur_lines = []
        continue

    # ANCRAGE sans crochets : "ANCRAGE IVOIRIEN"
    if re.match(r'^ANCRAGE\s+IVOIRIEN', s):
        flush()
        cur_type = "ANCRAGE"; cur_title = "ANCRAGE IVOIRIEN"; cur_lines = []
        continue

    # Titre de tableau juste avant les | (ex: "TABLEAU DES STRUCTURES...")
    if re.match(r'^TABLEAU\s+DES\s+', s):
        flush()
        cur_type = "TABLEAU"; cur_title = s.strip(); cur_lines = []
        continue

    # TABLEAU MARKDOWN — lignes | col | col |
    if re.match(r'^\|', s) and s.count('|') >= 2:
        if cur_type != "TABLEAU":
            flush()
            cur_type = "TABLEAU"; cur_title = ""; cur_lines = []
        # Ignorer lignes séparatrices |---|---|
        if not re.match(r'^\|[-:\s|]+\|', s):
            cur_lines.append(s)
        continue

    # Sous-sections Phase 2 : "A. TITRE", "B. TITRE" etc.
    if re.match(r'^[A-Z]\.\s+[A-ZÀÉÈÊÙÛ\(]', s):
        flush()
        add_sous_titre(s, color=C_BLEU_TITRE)
        continue

    # HARKNESS
    if re.match(r'^[▶►]\s*HARKNESS', s):
        flush()
        cur_type  = "HARKNESS"
        cur_title = s.lstrip("▶► ").strip()
        continue

    # À RETENIR — bordures ┌┐└┘ OU bloc FORMATION ... sans bordures
    if re.match(r'^[┌].*À RETENIR', s):
        flush()
        tc = re.sub(r'^[┌─\s]+', '', s).rstrip('─┐').strip()
        cur_type = "RETENIR"; cur_title = tc; continue
    if re.match(r'^[└]', s):
        flush(); continue
    # Bloc de règles sans bordure (ex: "FORMATION DU PRESENT CONTINUOUS")
    if re.match(r'^FORMATION\s+DU\s+', s):
        flush()
        cur_type  = "RETENIR"
        cur_title = s.strip()
        cur_lines = []
        continue

    # ATTENTION
    if re.match(r'^⚠', s):
        flush()
        cur_type  = "ATTENTION"
        cur_title = s.lstrip("⚠ ").strip()
        continue

    # TRACE ÉCRITE
    if re.match(r'^✎', s):
        flush(); cur_type = "TRACE"; cur_lines = [s]; continue

    # ANALOGIE — flush() systématique (gère 1 ou 2 analogies)
    if re.match(r'^~\s*ANALOGIE', s):
        flush()   # flush TRACE ou ANALOGIE précédente
        cur_type  = "ANALOGIE"
        cur_title = s.lstrip("~ ").strip()
        cur_lines = []
        continue

    # DÉFINITION DPFC
    if re.match(r'^◆\s*DÉFINITION', s):
        flush(); cur_type = "DEF"; cur_title = s.lstrip("◆ ").strip(); continue

    # EXERCICE GUIDÉ
    if re.match(r'^◉', s):
        flush(); cur_type = "EXOGUIDE"; cur_title = s.lstrip("◉ ").strip(); continue

    # EXERCICE ENTRAÎNEMENT
    if re.match(r'^◎', s):
        flush(); cur_type = "EXOENT"; cur_title = s.lstrip("◎ ").strip(); continue

    # DEVOIR MAISON
    if re.match(r'^◈', s):
        flush(); cur_type = "DEVOIR"; cur_title = s.lstrip("◈ ").strip(); continue

    # CORRIGÉ
    if re.match(r'^✔\s*CORRIGÉ', s):
        flush(); cur_type = "CORRIGE"; cur_title = s.lstrip("✔ ").strip(); continue

    # Bloc ════ final (pied de cours) → on arrête le contenu
    if re.match(r'^[═]{4,}', s) and i > body_start + 10:
        flush(); continue

    # Balise "CITATION FINALE" → déclenche le mode citation
    if re.match(r'^CITATION\s+FINALE', s):
        flush(); continue

    # Citation entre guillemets (n'importe laquelle, pas forcément la citation LPA)
    if re.match(r'^["""«]', s) and not cur_type:
        flush()
        m_en = re.search(r'["""«]([^"""»]+)["""»]', s)
        if m_en:
            citation_en = m_en.group(1).strip()
        continue
    # Traduction de la citation entre parenthèses
    if re.match(r'^\(', s) and citation_en and not cur_type:
        m_fr = re.search(r'\(([^)]+)\)', s)
        if m_fr:
            citation_fr = m_fr.group(1).strip()
        continue
    # Auteur de la citation "— Prénom Nom"
    if re.match(r'^[—–-]\s+\w', s) and citation_en and not cur_type:
        citation_sig = s.lstrip("—– ").strip()
        continue

    # Mots-clés LPA par défaut
    if "silence teaches nothing" in s:
        flush()
        m_en = re.search(r'"([^"]+)"', s)
        if m_en:
            citation_en = m_en.group(1)
        continue
    if "silence n'apprend rien" in s:
        m_fr = re.search(r'\(([^)]+)\)', s)
        if m_fr:
            citation_fr = m_fr.group(1)
        continue
    # Ligne signature finale Prof. … · DPFC/…
    if re.match(r'^Prof\.?\s+\w+', s) and ('DPFC' in s or 'Anglais' in s or '3ème' in s):
        citation_sig = s
        continue

    # Contenu d'un bloc ouvert
    if cur_type is not None:
        cur_lines.append(s)
    else:
        # Hors bloc
        if SOUS_TITRES.match(s):
            add_sous_titre(s)
        elif re.match(r'^(Titre\s*:|Compétence\s*:|Professeur\s+\w)', s):
            # Lignes de contexte Phase 1 — texte bleu gris léger
            add_text(s, color=C_GRIS_CLAIR, size=11, italic=True)
        elif re.match(r'^LIS CE TEXTE', s):
            # Instruction avant le texte de découverte — texte courant
            add_text(s, color=C_GRIS_TEXTE, size=11, sb=4)
        elif re.match(r'^TRADUCTION', s):
            # Entête "TRADUCTION COMPLÈTE :" hors boîte → ignorer
            pass
        else:
            add_text(s, color=C_GRIS_TEXTE, size=11)

flush()

# Citation finale
add_citation(citation_en, citation_fr,
             citation_sig or f"{professeur}  ·  {organisme}")

# ============================================================
#  SAUVEGARDE
# ============================================================
nom = "cours_lpa.docx"
if header_lines:
    nom = re.sub(r'[^\w\s·—-]', '', header_lines[0])
    nom = re.sub(r'[\s·—]+', '_', nom).strip('_') + ".docx"

doc.save(nom)
print(f"✅ Document généré : {nom}")
try:
    from google.colab import files
    files.download(nom)
except ImportError:
    print("(Hors Colab — fichier sauvegardé localement)")
